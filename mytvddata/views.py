# =========================
# DJANGO CORE IMPORTS
# =========================
import io
from io import BytesIO
from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse
from django.http import HttpResponse, JsonResponse, FileResponse, HttpResponseRedirect
from django.core.paginator import Paginator, PageNotAnInteger, EmptyPage
from django.db import transaction, IntegrityError
from django.db.models import (
    Count, Avg, Min, Max, Sum,
    CharField, Value, Q, Case, When, F, OuterRef, BooleanField
)
from django.utils import timezone
from django.db.models import Prefetch
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required, permission_required, user_passes_test
from django.views.decorators.csrf import csrf_exempt
from docx import Document
from django.db import transaction



# =========================
# DJANGO GENERIC VIEWS
# =========================
from django.views.generic import TemplateView, ListView, CreateView, DetailView, FormView
from django.views.generic.detail import SingleObjectMixin
 
# =========================
# DJANGO FORMS
# =========================
from django.forms.models import (
    inlineformset_factory,
    formset_factory,
    modelform_factory,
    modelformset_factory
)
 
# =========================
# PROJECT MODELS
# =========================
from .models import (
    Country, Component, Subcomponent, Indicator,
    RepositoryIndicator, SurveyProject, SurveyDataset,
    DocSave, ReportSave, LocationCountry, City,
    WarehouseScript, ApiData, StoreAPI,ApiExtract, 
    ApiFieldConfig,ApiFieldMapping,StaffMember
)



# =========================
# PROJECT FORMS
# =========================
from .forms import (
    CountryForm, IndicatorForm, IndicatorSearchForm,
    ComponentForm, Loadcomponent, SelectcomponentForm,
    ComponentsForm, SubcomponentsForm, SelectSubcomponentForm,
    DocSaveForm, ReportSaveForm, SurveyProjectForm,
    SelectUrlForm, SurveyDatasetForm, SelectSurveyForm,
    RepositoryIndicatorForm, WarehouseScriptForm,SelectUrlForms,
    S_table_nameForm, RepositoryUploadForm,LocationCountryForm,
    SelectRepositoryForm, UploadExcelForm,SurveyUploadForm,
    SelectAPIForm, StoreAPIForm,  ApiFieldSelectionForm,StaffMemberForm
)

# =========================
# EXTERNAL LIBRARIES
# =========================
import pandas as pd
from datetime import datetime
import json
import requests
import folium
import streamlit as st
from streamlit_folium import st_folium
import csv
import base64

from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet


from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm

# =========================
# EXCEL / DATA EXPORT
# =========================
from openpyxl import Workbook
from openpyxl.formatting.rule import ColorScaleRule
from openpyxl.styles import Color, PatternFill, Font, Border, Alignment, Side
from tablib import Dataset
from django.http import HttpResponse
from openpyxl.utils import get_column_letter
import openpyxl
from docx.shared import RGBColor
from openpyxl.chart import PieChart, BarChart, Reference

# =========================
# EMAIL / AUTH UTILITIES
# =========================
from django.core.mail import send_mail
from django.template.loader import render_to_string, get_template
from django.contrib.sites.shortcuts import get_current_site
from django.utils.http import urlsafe_base64_decode, urlsafe_base64_encode
from django.utils.encoding import force_bytes, force_str
from django.core.mail import send_mail
from django.contrib import messages
 
# Compatibilité ancienne version Django
try:
    from django.utils.encoding import force_text
except ImportError:
    from django.utils.encoding import force_str as force_text
 
# =========================
# PROJECT SETTINGS
# =========================
from TVDdata import settings
 
# =========================
# REST FRAMEWORK
# =========================
from rest_framework.response import Response
 
# =========================
# PROJECT RESOURCES
# =========================
from .resources import SurveyResource
import re




# =========================
# COVER PAGE WITH MAP
# =========================
def coverPage(request):
    """
    Affiche une carte Folium avec les localisations depuis la base de données
    """
     # Création de la carte centrée sur l’Afrique
    m = folium.Map(
        location=[0.0, 20.0],
        zoom_start=4,
        tiles="cartodb positron",
        scrollWheelZoom=True
    )

    folium.LayerControl().add_to(m)

    # ✅ Récupération des localisations
    locations = LocationCountry.objects.all()

    # ✅ Ajout des markers avec redirection immédiate
    for location in locations:
        if location.latitude and location.longitude:
            try:
                country = Country.objects.get(cca3=location.iso3)
                url = request.build_absolute_uri(
                    reverse('country_dashboard', args=[country.pk])
                )

                # Injection JS pour redirection au clic
                js = f"""
                    <script>
                        var marker = L.marker([{location.latitude}, {location.longitude}]).addTo({m.get_name()});
                        marker.bindTooltip("{location.name}");
                        marker.on('click', function() {{
                            window.location.href = '{url}';
                        }});
                    </script>
                """
                m.get_root().html.add_child(folium.Element(js))

                # Optionnel : popup plus grand avec bouton
                popup_html = f"""
                    <div style="width:250px; height:100px;">
                        <b>{location.name}</b><br>
                        <a href='{url}' class='btn btn-outline-primary btn-sm'>View profile</a>
                    </div>
                """
                folium.Marker(
                    location=(location.latitude, location.longitude),
                    popup=folium.Popup(popup_html, max_width=100),
                    tooltip=location.name#,
                   # icon=folium.Icon(color="blue", icon="globe", prefix="fa")
                ).add_to(m)

            except Country.DoesNotExist:
                folium.Marker(
                    location=(location.latitude, location.longitude),
                    tooltip=location.name,
                    icon=folium.Icon(color="gray", icon="globe", prefix="fa")
                ).add_to(m)

    # ✅ Conversion en HTML
    map_html = m._repr_html_()

    # Distincts pour statistiques
    all_countries = [c.name for c in Country.objects.all()]
    distinct_countries = set(all_countries)

    all_indicators = [i.indicator_code for i in Indicator.objects.all()]
    distinct_indicators = set(all_indicators)

    all_reportSave = [r.title_rep for r in ReportSave.objects.all()]
    distinct_reportSave = set(all_reportSave)
 
    context = {
        'm': map_html,
        'distinct_countries': len(distinct_countries),
        'distinct_indicators': len(distinct_indicators),
        'distinct_reportSave': len(distinct_reportSave)
    }
 
    return render(request, 'mytvddata/cover/index.html', context)


# =========================
# SIGNUP VIEW
# =========================
def signup(request):
    if request.method == "POST":
        username = request.POST.get('username')
        fname = request.POST.get('fname')
        lname = request.POST.get('lname')
        email = request.POST.get('email')
        pass1 = request.POST.get('pass1')
        pass2 = request.POST.get('pass2')
 
        # ✅ Vérifier si username existe
        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ Vérifier email existe déjà
        if User.objects.filter(email=email).exists():
            messages.error(request, "Email already registered")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ ✅ VALIDATION EMAIL (format)
        email_regex = r'^[\w\.-]+@[\w\.-]+\.\w+$'
        if not re.match(email_regex, email):
            messages.error(request, "Invalid email format")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ Longueur username
        if len(username) > 15:
            messages.error(request, "Username must be under 15 characters")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ Vérifier caractères username
        if not username.isalnum():
            messages.error(request, "Username must be alphanumeric")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ Vérifier mots de passe identiques
        if pass1 != pass2:
            messages.error(request, "Passwords do not match")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ ✅ SÉCURITÉ MOT DE PASSE
        if len(pass1) < 8:
            messages.error(request, "Password must be at least 8 characters long")
            return render(request, 'mytvddata/cover/signup.html')
 
        if not re.search(r'[A-Z]', pass1):
            messages.error(request, "Password must contain at least one uppercase letter")
            return render(request, 'mytvddata/cover/signup.html')
 
        if not re.search(r'[a-z]', pass1):
            messages.error(request, "Password must contain at least one lowercase letter")
            return render(request, 'mytvddata/cover/signup.html')
 
        if not re.search(r'[0-9]', pass1):
            messages.error(request, "Password must contain at least one number")
            return render(request, 'mytvddata/cover/signup.html')
 
        # ✅ Création utilisateur
        user = User.objects.create_user(
            username=username,
            email=email,
            password=pass1
        )
 
        user.first_name = fname
        user.last_name = lname
        user.is_active = True
 
        user.save()
 
        messages.success(request, "Account created successfully")
        return redirect('signin')
 
    return render(request, 'mytvddata/cover/signup.html')
 
 
# =========================
# SIGNIN VIEW
# =========================
def signin(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('pass1')
 
        user = authenticate(request, username=username, password=password)
 
        if user is not None:
            login(request, user)
            messages.success(request, "Logged in successfully")
            return redirect('coverPage')
        else:
            messages.error(request, "Invalid username or password")
            return render(request, 'mytvddata/cover/signin.html')
 
    return render(request, 'mytvddata/cover/signin.html')
 
# =========================
# SIGNOUT VIEW
# =========================
@login_required
def signout(request):
    logout(request)
    messages.success(request, "Logged out successfully")
    return redirect('coverPage')
 
 
def countries(request):
    country_list = Country.objects.order_by('name')
 
    query = request.GET.get('q')
    if query:
        country_list = country_list.filter(name__icontains=query)
 
    paginator = Paginator(country_list, 10)
    countries_page = paginator.get_page(request.GET.get('page'))
 
    return render(request, 'mytvddata/pages/country/countries.html', {
        'country_list': countries_page,
        'total_country': paginator.count
    })


 
# =========================
# CRUD COUNTRY (PRO VERSION)
# =========================
@login_required
def add_country(request):
 
    # ✅ Base queryset optimisé
    country_list = Country.objects.all().order_by('name')
 
    # =========================
    # SEARCH
    # =========================
    search_query = request.GET.get('name')
    if search_query:
        country_list = country_list.filter(name__icontains=search_query)
 
    # =========================
    # PAGINATION (VERSION PRO)
    # =========================
    paginator = Paginator(country_list, 50)
    countries_page = paginator.get_page(request.GET.get('page'))
 
    # =========================
    # FORM INIT
    # =========================
    form = CountryForm()
 
    # =========================
    # POST ACTIONS (CRUD)
    # =========================
    if request.method == 'POST':
 
        # ✅ SAVE / UPDATE
        if 'save' in request.POST:
            pk = request.POST.get('save')
 
            if pk:  # update
                country = get_object_or_404(Country, id=pk)
                form = CountryForm(request.POST, instance=country)
            else:  # create
                form = CountryForm(request.POST)
 
            if form.is_valid():
                form.save()
                messages.success(request, "Country saved successfully")
                return redirect('add_country')  # ✅ éviter resubmit
 
        # ✅ DELETE
        elif 'delete' in request.POST:
            if request.user.groups.filter(name='Manager and delete').exists():
 
                pk = request.POST.get('delete')
                country = get_object_or_404(Country, id=pk)
                country.delete()
 
                messages.success(request, "Country deleted successfully")
                return redirect('add_country')
 
            else:
                return render(request, 'pages/examples/403.html')
 
        # ✅ EDIT (pré-remplir form)
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            country = get_object_or_404(Country, id=pk)
            form = CountryForm(instance=country)
 
    # =========================
    # CONTEXT
    # =========================
    context = {
        'countries': countries_page,
        'total_country': paginator.count,
        'form': form,
        'search_query': search_query
    }
 
    return render(request, 'mytvddata/pages/country/add_country.html', context)
 

# =========================
# EXPORT COUNTRY PRO MAX
# =========================
@login_required
def export_country(request):
 
    # ✅ Réponse HTTP
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=Country_Report.xlsx'
 
    wb = Workbook()
 
    # =========================
    # SHEET 1: ALL COUNTRIES
    # =========================
    ws = wb.active
    ws.title = "Countries"
 
    headers = [
        "ID", "Name", "CCA3", "Capital",
        "Subregion", "Population"
    ]
 
    ws.append(headers)
 
    # ✅ STYLE HEADER
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    header_font = Font(bold=True)
 
    for col_num, col_title in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
 
    # =========================
    # FILTRE DYNAMIQUE
    # =========================
    query = request.GET.get('name')
    countries = Country.objects.all()
 
    if query:
        countries = countries.filter(name__icontains=query)
 
    # =========================
    # DATA
    # =========================
    for c in countries:
        ws.append([
            c.id,
            c.name,
            c.cca3,
            c.capital,
            c.subregion,
            c.population
        ])
 
    # ✅ FILTER EXCEL
    ws.auto_filter.ref = ws.dimensions
 
    # ✅ AUTO WIDTH
    for col in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
 
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
 
        ws.column_dimensions[col_letter].width = max_length + 2
 
    # =========================
    # SHEET 2: SUMMARY
    # =========================
    ws2 = wb.create_sheet(title="Summary")
 
    ws2.append(["Metric", "Value"])
    ws2.append(["Total Countries", countries.count()])
 
    # Exemple métrique supplémentaire
    total_pop = sum([c.population or 0 for c in countries])
    ws2.append(["Total Population", total_pop])
 
    # =========================
    # SAVE
    # =========================
    wb.save(response)
 
    return response
 



@login_required
def list_indicator(request):

    indicators_qs = Indicator.objects.select_related('subcomponent').order_by('indicator_name')

    form_subcomponent = SelectSubcomponentForm(request.GET or None)

    search_name = request.GET.get('indicator_name')
    by_subcomponent = request.GET.get('by_subcomponent')

    if search_name:
        indicators_qs = indicators_qs.filter(indicator_name__icontains=search_name)

    if by_subcomponent:
        indicators_qs = indicators_qs.filter(subcomponent__id=by_subcomponent)

    paginator = Paginator(indicators_qs, 10)
    indicators_page = paginator.get_page(request.GET.get('page'))

    # ✅ total par maladie
    total_by_disease = indicators_qs.values('category_indicator').annotate(total=Count('id'))

    context = {
        'indicators': indicators_page,
        'total_indicator': paginator.count,
        'form_subcomponent': form_subcomponent,
        'by_subcomponent': by_subcomponent,
        'total_by_disease': total_by_disease,
    }

    return render(request, 'mytvddata/pages/indicator/indicator_list.html', context)


def save_indicator_form(request, form, template_name):

    data = {}
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            data['form_is_valid'] = True

            indicators = Indicator.objects.select_related('subcomponent')

            data['html_indicator_list'] = render_to_string(
                'mytvddata/pages/indicator/partial_indicator_list.html',
                {'indicators': indicators},
                request=request
            )
        else:
            data['form_is_valid'] = False

    context = {'form': form}
    data['html_form'] = render_to_string(template_name, context, request=request)

    return JsonResponse(data)


@login_required
def create_indicator(request):

    if request.method == 'POST':
        form = IndicatorForm(request.POST)

    else:
        form = IndicatorForm()
    return save_indicator_form(request, form, 'mytvddata/pages/indicator/partial_indicator_create.html')




@login_required
def update_indicator(request, pk):

    indicator = get_object_or_404(Indicator, pk=pk)
    data = dict()

    if request.method == 'POST':
        form = IndicatorForm(request.POST, instance=indicator)
        if form.is_valid():
            form.save()
            data['form_is_valid'] = True
            indicators = Indicator.objects.select_related('subcomponent')
            data['html_indicator_list'] = render_to_string(
                'mytvddata/pages/indicator/partial_indicator_list.html',
                {'indicators': indicators},
                request=request
            )
        else:
            data['form_is_valid'] = False
    else:
        form = IndicatorForm(instance=indicator)

    context = {'form': form, 'indicator': indicator}
    data['html_form'] = render_to_string(
        'mytvddata/pages/indicator/partial_indicator_update.html',
        context,
        request=request
    )
    return JsonResponse(data)


@login_required
def delete_indicator(request, pk):

    indicator = get_object_or_404(Indicator, pk=pk)
    data = dict()

    if request.method == 'POST':
        indicator.delete()
        data['form_is_valid'] = True

        indicators = Indicator.objects.select_related('subcomponent')

        data['html_indicator_list'] = render_to_string(
            'mytvddata/pages/indicator/partial_indicator_list.html',
            {'indicators': indicators},
            request=request
        )

    else:
        context = {'indicator': indicator}
        data['html_form'] = render_to_string(
            'mytvddata/pages/indicator/partial_indicator_delete.html',
            context,
            request=request
        )

    return JsonResponse(data)





@login_required
def export_indicator(request):
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=Indicators.xlsx'

    wb = Workbook()
    ws = wb.active
    ws.title = "Indicators"

    # American institutional style headers
    headers = [
        "ID", "Indicator Code", "Indicator Name", "Description",
        "Target Value", "Metric", "Unit of Measure",
        "Reporting Frequency", "Indicator Type", "Data Source",
        "Health Domain", "Category",
        "Forecasting", "Performance", "Reference Data"
    ]
    ws.append(headers)

    # Style headers: bold, centered, blue background, white text
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")

    for col_num, column_title in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    indicators = Indicator.objects.select_related('subcomponent')

    for i in indicators:
        ws.append([
            i.id,
            i.indicator_code,
            i.indicator_name,
            i.indicator_description,
            i.indicator_target,
            i.indicator_metric,
            i.indicator_unit,
            i.indicator_frequency,
            i.type_indicator,
            i.indicator_source,
            i.subcomponent.subcomponent_name if i.subcomponent else '',
            i.category_indicator,
            i.forecasting_indicator,
            i.performance_indicator,
            i.ref_data
        ])

    # Auto-fit column widths
    for col in ws.columns:
        max_length = 0
        column = col[0].column  # Get column index
        column_letter = get_column_letter(column)
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column_letter].width = adjusted_width

    wb.save(response)
    return response





@login_required
def import_indicator(request):

    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active

        # ✅ Vérification des en-têtes
        expected_headers = [
            "ID", "Code", "Name", "Description",
            "Target", "Metric", "Unit",
            "Frequency", "Type", "Source",
            "Subcomponent", "Category",
            "Forecast", "Performance", "Ref Data"
        ]
        headers = [cell.value for cell in ws[1]]
        if headers != expected_headers:
            messages.error(request, "Invalid Excel format. Headers must match the template.")
            return redirect('import_indicator')

        # ✅ Vérification des choix autorisés
        valid_categories = dict(Indicator.CATOGORY_CHOICES).keys()
        valid_forecasts = dict(Indicator.MODELING_CHOICES).keys()
        valid_performance = dict(Indicator.PERFORMANCE_CHOICES).keys()
        valid_types = dict(Indicator.TYPE_CHOICES).keys()

        errors = []
        for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            code, name, desc, target, metric, unit, freq, type_ind, source, subcomp, category, forecast, performance, ref_data = (
                row[1], row[2], row[3], row[4], row[5], row[6], row[7],
                row[8], row[9], row[10], row[11], row[12], row[13], row[14]
            )

            # Vérification des champs obligatoires
            if not code or not name:
                errors.append(f"Row {idx}: Code and Name are required.")

            # Vérification des choix
            if category and category not in valid_categories:
                errors.append(f"Row {idx}: Invalid category '{category}'. Must be one of {list(valid_categories)}.")

            if forecast and forecast not in valid_forecasts:
                errors.append(f"Row {idx}: Invalid forecast '{forecast}'. Must be one of {list(valid_forecasts)}.")

            if performance and performance not in valid_performance:
                errors.append(f"Row {idx}: Invalid performance '{performance}'. Must be one of {list(valid_performance)}.")

            if type_ind and type_ind not in valid_types:
                errors.append(f"Row {idx}: Invalid type '{type_ind}'. Must be one of {list(valid_types)}.")

        # ✅ Si erreurs, on arrête
        if errors:
            for e in errors:
                messages.error(request, e)
            return redirect('import_indicator')

        # ✅ Sinon, insertion en base
        for row in ws.iter_rows(min_row=2, values_only=True):
            Indicator.objects.create(
                indicator_code=row[1],
                indicator_name=row[2],
                indicator_description=row[3],
                indicator_target=row[4],
                indicator_metric=row[5],
                indicator_unit=row[6],
                indicator_frequency=row[7],
                type_indicator=row[8],
                indicator_source=row[9],
                category_indicator=row[11],
                forecasting_indicator=row[12],
                performance_indicator=row[13],
                ref_data=row[14],
            )

        messages.success(request, 'Indicators imported successfully')
        return redirect('list_indicator')

    return render(request, 'mytvddata/pages/indicator/import_indicator.html')







@login_required
def import_preview(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        excel_file = request.FILES['excel_file']
        wb = openpyxl.load_workbook(excel_file)
        ws = wb.active

        # Vérification des en-têtes
        expected_headers = [
            "ID", "Code", "Name", "Description",
            "Target", "Metric", "Unit",
            "Frequency", "Type", "Source",
            "Subcomponent", "Category",
            "Forecast", "Performance", "Ref Data"
        ]
        headers = [cell.value for cell in ws[1]]
        if headers != expected_headers:
            messages.error(request, "Invalid Excel format. Headers must match the template.")
            return redirect('list_indicator')

        preview_data = []
        for idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            preview_data.append({
                "row": idx,
                "code": row[1],
                "name": row[2],
                "description": row[3],
                "target": row[4],
                "metric": row[5],
                "unit": row[6],
                "frequency": row[7],
                "type": row[8],
                "source": row[9],
                "subcomponent": row[10],
                "category": row[11],
                "forecast": row[12],
                "performance": row[13],
                "ref_data": row[14],
            })

             # ✅ Nettoyer avant de stocker
        if 'import_content' in request.session:
            del request.session['import_content']

        # ✅ Stocker le contenu encodé en Base64
        excel_file.seek(0)
        encoded_content = base64.b64encode(excel_file.read()).decode('utf-8')
        request.session['import_content'] = encoded_content

        return render(request, 'mytvddata/pages/indicator/import_preview.html', {
            'preview_data': preview_data,
        })

    return redirect('list_indicator')





@login_required
def confirm_import(request):
    if request.method == 'POST':
        encoded_content = request.session.get('import_content')
        if not encoded_content:
            messages.error(request, "No file found in session.")
            return redirect('list_indicator')

        # ✅ Décoder le contenu Base64 en bytes
        file_content = base64.b64decode(encoded_content)
        wb = openpyxl.load_workbook(io.BytesIO(file_content))
        ws = wb.active

        for row in ws.iter_rows(min_row=2, values_only=True):
            Indicator.objects.create(
                indicator_code=row[1],
                indicator_name=row[2],
                indicator_description=row[3],
                indicator_target=row[4],
                indicator_metric=row[5],
                indicator_unit=row[6],
                indicator_frequency=row[7],
                type_indicator=row[8],
                indicator_source=row[9],
                category_indicator=row[11],
                forecasting_indicator=row[12],
                performance_indicator=row[13],
                ref_data=row[14],
            )

 # ✅ Nettoyer après import
        if 'import_content' in request.session:
            del request.session['import_content']

        messages.success(request, "Indicators imported successfully!")
        return redirect('list_indicator')

    # ✅ Nettoyer si POST invalide
    if 'import_content' in request.session:
        del request.session['import_content']

    return redirect('list_indicator')





# Télécharger modèle Excel
def download_indicator_template(request):
    df = pd.DataFrame(columns=[
        "ID", "Code", "Name", "Description",
            "Target", "Metric", "Unit",
            "Frequency", "Type", "Source",
            "Subcomponent", "Category",
            "Forecast", "Performance", "Ref Data"
    ])
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="indicator_template.xlsx"'
    df.to_excel(response, index=False)
    return response



# CREATE VIEWS FOR SCRIPTS TABLE------------------------------------

def t_scripts_view(request):
  
    form_survey_title = S_table_nameForm(request.GET or None)
    by_title =  request.GET.get('by_title')
    end_date =  request.GET.get('end')
  
    t_scripts_views = WarehouseScript.objects.all().order_by("ws_title")
    form = WarehouseScriptForm()
    if request.method=='GET':
        reseash_s_focal_point=request.GET.get('ws_focalpoint')
        if reseash_s_focal_point!=None:
            t_scripts_views= WarehouseScript.objects.filter(ws_focalpoint__icontains=reseash_s_focal_point)
        if by_title!=None:
            t_scripts_views = WarehouseScript.objects.all().filter(Q(id__icontains=by_title),Q(ws_date__lte=end_date)).order_by("ws_title")# id use to filter the title because the paramters is id but not label
               
    context={'form_survey_title': form_survey_title,
             'form':form,
             't_scripts_views': t_scripts_views,
             'by_title':by_title,
             'end_date':end_date
             }
        
    return render(request, 'mytvddata/pages/scripts_view/t_scripts_view.html', context)

# Method for delete SCRIPTS

def delete_t_scripts(request, pk):
    if request.method == "POST":
        WarehouseScript.objects.filter(id=pk).delete()
        return JsonResponse({'msg': 'deleted'})
    else:
        t_script = WarehouseScript.objects.get(id=pk)
        return render(request, 'mytvddata/pages/scripts_view/delete_t_scripts_modal.html',
                      {'t_script': t_script})


# Method for add new SCRIPTS

def add_new_t_scripts(request):
    if request.method == "POST":
        form = WarehouseScriptForm(request.POST)
        if form.is_valid():
            form.save(commit=True)
          #  form.save_m2m
         #   return JsonResponse({'msg': 'Data saved'})
            messages.success(request, "SCRIPTS Added Successfully")
            return redirect('t_scripts_view')
        else:
            print("ERROR FORM INVALID")
            return JsonResponse({'msg': 'ERROR FORM INVALID'})
    else:
   #     form = WorkplanForm()
  #  return JsonResponse({'form': form})
        context = {
            'form': form
        }
    return render(request, 'mytvddata/pages/scripts_view/add_t_scripts_modal.html', context)

         
# Method for edit SCRIPTS
def edit_t_scripts(request, pk):
    t_script = WarehouseScript.objects.get(id=pk)
    if request.method == "POST":
        form = WarehouseScriptForm(request.POST, instance=t_script)
        if form.is_valid():
            form.save()
            return JsonResponse({'msg': 'updated'})
        return JsonResponse({'msg': 'error'})
    else:
        form = WarehouseScriptForm(instance=t_script)
        return render(request, 'mytvddata/pages/scripts_view/edit_t_scripts_modal.html',
                      {'form': form, 't_script': t_script})



# Method for import and export Table SurveyDataset


# Gestion des projets Survey (CRUD + pagination + recherche)
def survey_add_project(request):
    form = SurveyProjectForm()
    surveyProjects = SurveyProject.objects.all()

    # 🔎 Recherche par titre
    search_title = request.GET.get('title_surv')
    if search_title:
        surveyProjects = surveyProjects.filter(title_surv__icontains=search_title)

    # 📑 Pagination
    paginator = Paginator(surveyProjects, 3)
    page = request.GET.get('page')
    try:
        surveyProjects = paginator.page(page)
    except PageNotAnInteger:
        surveyProjects = paginator.page(1)
    except EmptyPage:
        surveyProjects = paginator.page(paginator.num_pages)

    # 📝 Gestion des actions POST

    
    # Actions POST
    if request.method == "POST":
        if "save" in request.POST:
            pk = request.POST.get("save")

            if pk:
                surveyProject = get_object_or_404(SurveyProject, id=pk)
                form = SurveyProjectForm(request.POST, instance=surveyProject)
            else:
                form = SurveyProjectForm(request.POST)


            if form.is_valid():
                project = form.save(commit=False)
                # ✅ Forcer une date par défaut si vide
                if not project.start_date:
                    project.start_date = timezone.now().date()
                if not project.end_date:
                    project.end_date = project.start_date  # ou une autre logique
                project.save()


              #  form.save()
                messages.success(request, '✅ Donnée enregistrée ou mise à jour')
                return redirect('project_survey')

        elif "delete" in request.POST:
            pk = request.POST.get('delete')
            surveyProject = get_object_or_404(SurveyProject, id=pk)
            surveyProject.delete()
            messages.success(request, 'Survey project supprimé avec succès')
            return redirect('project_survey')

        elif "edit" in request.POST:
            pk = request.POST.get('edit')
            surveyProject = get_object_or_404(SurveyProject, id=pk)
            form = SurveyProjectForm(instance=surveyProject)
        
           
    context = {
        'surveyProjects': surveyProjects,
        'paginator_survey': paginator,
        'form': form,
        'title': 'home',
    }
    return render(request, 'mytvddata/pages/survey/survey_project.html', context)









# Télécharger modèle Excel
def download_survey_template(request):
    df = pd.DataFrame(columns=[
        "surveyProject_id", "quest_code", "question",
        "response_text", "response_num", "level_1", "level_2"
    ])
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="survey_template.xlsx"'
    df.to_excel(response, index=False)
    return response


# Prévisualiser import
def preview_import_survey(request):
    preview_data = []
    if request.method == "POST" and request.FILES.get("excel_file"):
        excel_file = request.FILES["excel_file"]

        try:
            df = pd.read_excel(excel_file)

            # 🔑 Conversion sécurisée de toutes les valeurs
            normalized_data = []
            for row in df.to_dict(orient="records"):
                clean_row = {}
                for key, value in row.items():
                    if isinstance(value, pd.Timestamp):
                        clean_row[key] = value.date().isoformat()
                    elif isinstance(value, (pd.Series, pd.DataFrame)):
                        clean_row[key] = str(value)
                    elif pd.isna(value):
                        clean_row[key] = None
                    else:
                        clean_row[key] = str(value)
                normalized_data.append(clean_row)

            preview_data = normalized_data
            request.session["import_data"] = preview_data
            messages.info(request, f"{len(preview_data)} lignes détectées pour importation.")

        except Exception as e:
            messages.error(request, f"Erreur lors de la lecture du fichier : {e}")

    return render(request, "mytvddata/pages/survey/survey_dataset.html", {
        "preview_data": preview_data,
        "surveyDatasets": SurveyDataset.objects.all().order_by("-surveyProject__end_date"),
    })


# Confirmer import
def confirm_import_survey(request):
    import_report = {"total": 0, "created": 0, "updated": 0, "skipped": 0, "errors": 0}
    data = request.session.get("import_data", [])
    report_data = []

    for row in data:
        import_report["total"] += 1
        try:
            # Vérifier si une ligne identique existe déjà
            existing = SurveyDataset.objects.filter(
                surveyProject_id=row.get("surveyProject_id"),
                quest_code=row.get("quest_code"),
                question=row.get("question"),
                response_text=row.get("response_text"),
                response_num=row.get("response_num"),
                level_1=row.get("level_1"),
                level_2=row.get("level_2"),
            ).first()

            if existing:
                # Vérifier si des champs diffèrent
                changed = False
                for field in ["surveyProject_id", "quest_code", "question", "response_text", "response_num", "level_1", "level_2"]:
                    new_val = row.get(field)
                    if getattr(existing, field) != new_val:
                        setattr(existing, field, new_val)
                        changed = True

                if changed:
                    existing.save()
                    import_report["updated"] += 1
                    report_data.append({"status": "Mis à jour", "raison": "Valeurs différentes", **row})
                else:
                    import_report["skipped"] += 1
                    report_data.append({"status": "Ignoré", "raison": "Valeurs identiques", **row})
            else:
                # Création si nouvelle ligne
                SurveyDataset.objects.create(
                    surveyProject_id=row.get("surveyProject_id"),
                    quest_code=row.get("quest_code"),
                    question=row.get("question"),
                    response_text=row.get("response_text"),
                    response_num=row.get("response_num"),
                    level_1=row.get("level_1"),
                    level_2=row.get("level_2"),
                )
                import_report["created"] += 1
                report_data.append({"status": "Créé", "raison": "Nouvelle ligne", **row})

        except Exception as e:
            import_report["errors"] += 1
            report_data.append({"status": "Erreur", "raison": str(e), **row})

    # Sauvegarder rapport en session pour export Excel/PDF
    request.session["last_import_report"] = import_report
    request.session["last_import_report_data"] = report_data
    request.session["import_data"] = []

    messages.success(
        request,
        f"Import terminé : {import_report['created']} créés, {import_report['updated']} mis à jour, "
        f"{import_report['skipped']} ignorés, {import_report['errors']} erreurs."
    )

    return render(request, "mytvddata/pages/survey/survey_dataset.html", {
        "import_report": import_report,
        "report_data": report_data,  # pour affichage détaillé
        "surveyDatasets": SurveyDataset.objects.all().order_by("-surveyProject__end_date"),
    })








# Export rapport en Excel
def export_surveyimport_report(request):
    report = request.session.get("last_import_report", {
        "total": 0, "created": 0, "updated": 0, "skipped": 0, "errors": 0
    })
    df = pd.DataFrame([report])

    output = BytesIO()
    df.to_excel(output, index=False)
    output.seek(0)

    response = HttpResponse(output, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="import_report.xlsx"'
    return response





def export_import_report_pdf(request):
    report = request.session.get("last_import_report", {
        "total": 0, "created": 0, "updated": 0, "skipped": 0, "errors": 0
    })
    report_data = request.session.get("last_import_report_data", [])

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="import_report.pdf"'

    doc = SimpleDocTemplate(response, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()

    # En-tête
    elements.append(Paragraph("Rapport d'importation SurveyDataset", styles["Title"]))
    elements.append(Spacer(1, 12))

    # Résumé
    summary_text = (
        f"Total lignes traitées : {report['total']}<br/>"
        f"Lignes créées : {report['created']}<br/>"
        f"Lignes mises à jour : {report['updated']}<br/>"
        f"Lignes ignorées : {report['skipped']}<br/>"
        f"Erreurs : {report['errors']}"
    )
    elements.append(Paragraph(summary_text, styles["Normal"]))
    elements.append(Spacer(1, 24))

    # Tableau détaillé
    if report_data:
        data = [["Status", "Raison", "Survey Project", "Code Question", "Question",
                 "Réponse (Texte)", "Réponse (Numérique)", "Zone 1", "Zone 2"]]

        for row in report_data:
            data.append([
                row.get("status", ""),
                row.get("raison", ""),
                row.get("surveyProject_id", ""),
                row.get("quest_code", ""),
                row.get("question", ""),
                row.get("response_text", ""),
                row.get("response_num", ""),
                row.get("level_1", ""),
                row.get("level_2", ""),
            ])

        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(table)

    doc.build(elements)
    return response


# Importation des données SurveyDataset depuis Excel
def simple_upload(request):
    if request.method == 'POST':
        survey_resource = SurveyResource()
        dataset = Dataset()
        new_survey = request.FILES.get('myfile')

        if not new_survey:
            messages.error(request, 'Veuillez sélectionner un fichier')
            return redirect('simple_upload')

        if not new_survey.name.endswith('.xlsx'):
            messages.error(request, 'Format incorrect, utilisez un fichier Excel (.xlsx)')
            return redirect('simple_upload')

        try:
            df = pd.read_excel(new_survey)
            for _, row in df.iterrows():
                SurveyDataset.objects.create(
                    surveyProject_id=row.get("surveyProject_id"),  # 🔗 FK vers projet
                    quest_code=row.get("quest_code"),
                    question=row.get("question"),
                    response_text=row.get("response_text"),
                    response_num=row.get("response_num"),
                    level_1=row.get("level_1"),
                    level_2=row.get("level_2"),
                )
            messages.success(request, "Importation réussie !")
        except Exception as e:
            messages.error(request, f"Erreur lors de l'importation : {e}")

    return render(request, 'mytvddata/pages/survey/upload.html')




# Ajouter / modifier / supprimer des données SurveyDataset
def survey_add_data(request):
    context = {}
    form_survey_title = SelectSurveyForm(request.GET or None)
    by_survey = request.GET.get('by_survey')
    end_day = request.GET.get('end')

    # 🔎 Recherche et filtrage
    surveyDatasets = SurveyDataset.objects.all().order_by("-surveyProject__end_date")

    quest_code = request.GET.get('quest_code')
    if quest_code:
        surveyDatasets = surveyDatasets.filter(
            quest_code__icontains=quest_code
        ).order_by("-surveyProject__end_date")

    if by_survey and end_day:
        surveyDatasets = surveyDatasets.filter(
            Q(surveyProject__id__icontains=by_survey) &
            Q(surveyProject__end_date__lte=end_day)
        ).order_by("quest_code")

    # 📝 Formulaire par défaut
    form = SurveyDatasetForm()

    # ⚡ Gestion des actions POST
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if pk:
                surveyDataset = get_object_or_404(SurveyDataset, id=pk)
                form = SurveyDatasetForm(request.POST, instance=surveyDataset)
            else:
                form = SurveyDatasetForm(request.POST)

            if form.is_valid():
                form.save()
                messages.success(request, 'Survey dataset enregistré ou mis à jour avec succès')
                return redirect('index_survey')

        elif 'delete' in request.POST:
          #  if request.user.groups.filter(name='Manager and delete').exists():
                pk = request.POST.get('delete')
                surveyDataset = get_object_or_404(SurveyDataset, id=pk)
                surveyDataset.delete()
                messages.success(request, 'Survey dataset supprimé avec succès')
                return redirect('index_survey')
          # else:
          #      return render(request, 'mytvddata/cover/403.html')

        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            surveyDataset = get_object_or_404(SurveyDataset, id=pk)
            form = SurveyDatasetForm(instance=surveyDataset)

    # 📦 Contexte
    context.update({
        'surveyDatasets': surveyDatasets,
        'form': form,
        'form_survey_title': form_survey_title,
        'by_survey': by_survey,
        'end_day': end_day,
        'title': 'home',
    })

    return render(request, 'mytvddata/pages/survey/survey_dataset.html', context)



# Method for list all Survey dataset submitted
#@permission_required('myuhp.view_surveydataset',raise_exception=True)

def survey_report(request):
    surveyProjects = SurveyProject.objects.all().order_by("title_surv") 
    if request.method=='GET':
        st=request.GET.get('responsible')
        if st!=None:
            surveyProjects= SurveyProject.objects.filter(responsible__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(surveyProjects, num_of_items)
       
    try:
        surveyProjects = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        surveyProjects = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        surveyProjects = paginator.page(page)
            
    data={'surveyProjects': surveyProjects, 'paginator':paginator}
    
    return render(request, 'mytvddata/pages/survey/survey_report.html', data)

# Method for view single Survey DataSet page

def single_survey_page(request, pk):
    try:
        surveyProject=SurveyProject.objects.get(id=pk) 
        surveyDatasets =SurveyProject.objects.filter(Q(id__icontains=pk) & Q(project_surveyData__quest_code__isnull=False)).values(
        'responsible',
        'title_surv',
        'start_date',
        'end_date',
        'location_survey',
        'project_surveyData__quest_code',
        'project_surveyData__question',
        'project_surveyData__response_text',
        'project_surveyData__response_num',
        'project_surveyData__level_1',
        'project_surveyData__level_2'
    ).order_by('-start_date')
        
        
        page = request.GET.get('page')
        num_of_items = 10
        # paginator of country 
        paginator = Paginator(surveyDatasets, num_of_items)
        
        try:
            surveyDatasets = paginator.page(page)
        except PageNotAnInteger:
            page = 1
            surveyDatasets = paginator.page(page)
        except EmptyPage:
            page = paginator.num_pages
            surveyDatasets = paginator.page(page)
        
        data={'surveyProject':surveyProject, 'surveyDatasets':surveyDatasets, 'paginator':paginator}   
    except:
        data={'message':'The Survey DataSet you requested doesn t\' exist or isn t in the submitted results'}
    
    return render(request,'mytvddata/pages/survey/single_survey_page.html',data)


#Method to export ALL SURVEY PROJECT on Excel file
def export_project_survey(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Project_survey.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Project Survey"

    # Add headers
    headers = ["id","responsible", "title_surv", "start_date", "end_date", "location_survey"]
    ws.append(headers)

    # Add data from the model
    surveyProjects = SurveyProject.objects.all()
    for surveyProject in surveyProjects:
        ws.append([surveyProject.pk, surveyProject.responsible, surveyProject.title_surv, surveyProject.start_date, surveyProject.end_date, surveyProject.location_survey])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

#Method to export SURVEY DATASET ALL DATA on Excel file
def export_to_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="survey.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Survey"

    # Add headers
    headers = ["responsible", "title_surv", "start_date", "end_date", "location_survey", "quest_code", "question", "response_text", "response_num","level_1","level_2"]
    ws.append(headers)

    # Add data from the model
    surveyDatasets = SurveyDataset.objects.all()
    for s in surveyDatasets:
        ws.append([s.surveyProject.responsible, s.surveyProject.title_surv, s.surveyProject.start_date, s.surveyProject.end_date, s.surveyProject.location_survey, 
                   s.quest_code,  s.question,  s.response_text, s.response_num, s.level_1, s.level_2])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

#Method to export DATASET on Excel file ONE BY ONE PROJECT




 
 
# =========================
# EXPORT SURVEY PAGE (PRO)
# =========================
@login_required
def export_survey_page(request, pk):
 
    # ✅ Réponse HTTP moderne (.xlsx)
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=survey_dataset.xlsx'
 
    wb = Workbook()
    ws = wb.active
    ws.title = "Survey"
 
    # =========================
    # HEADERS
    # =========================
    headers = [
        'Responsible', 'Title', 'Start Date', 'End Date',
        'Location', 'Question Code', 'Question',
        'Response Text', 'Response Num', 'Level 1', 'Level 2'
    ]
 
    ws.append(headers)
 
    # =========================
    # QUERY OPTIMISÉE
    # =========================
    projects = SurveyProject.objects.filter(pk=pk).prefetch_related(
        Prefetch(
            'project_surveyData',
            queryset=SurveyDataset.objects.filter(quest_code__isnull=False)
        )
    )
 
    # =========================
    # DATA
    # =========================
    for project in projects:
        for data in project.project_surveyData.all():
            ws.append([
                project.responsible,
                project.title_surv,
                project.start_date,
                project.end_date,
                project.location_survey,
                data.quest_code,
                data.question,
                data.response_text,
                data.response_num,
                data.level_1,
                data.level_2
            ])
 
    # =========================
    # AUTO WIDTH (UX)
    # =========================
    for col in ws.columns:
        max_length = 0
        col_letter = col[0].column_letter
 
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
 
        ws.column_dimensions[col_letter].width = max_length + 2
 
    # =========================
    # SAVE
    # =========================
    wb.save(response)
 
    return response
 

              
              
## UPLOAD DATASET OF REPOSITORY FORM EXCEL FILE



# 📥 Upload et import Excel
def upload_repository(request):
    preview_data = None
    if request.method == 'POST':
        new_repository = request.FILES.get('myfile')
        if not new_repository or not new_repository.name.endswith('.xlsx'):
            messages.error(request, 'Format incorrect, veuillez uploader un fichier .xlsx')
            return render(request, 'mytvddata/pages/repository/upload_repository.html')

        # Prévisualisation avec pandas
        df = pd.read_excel(new_repository)
        preview_data = df.head(5).to_html(classes="table table-striped")

        # Import dans la base
        for _, row in df.iterrows():
            RepositoryIndicator.objects.create(
                indicator_code=row.get("indicator_code"),
                country=row.get("country"),
                spatial_dim=row.get("spatial_dim"),
                indicator_name=row.get("indicator_name"),
                dim1_type=row.get("dim1_type"),
                dim1=row.get("dim1"),
                dim2_type=row.get("dim2_type"),
                dim2=row.get("dim2"),
                dim3_type=row.get("dim3_type"),
                dim3=row.get("dim3"),
                time_dim=row.get("time_dim"),
                alpha_value=row.get("alpha_value"),
                numeric_value=row.get("numeric_value"),
                publish_date=row.get("publish_date"),
            )
        messages.success(request, "Données importées avec succès ✅")

    return render(request, "mytvddata/pages/repository/upload_repository.html", {
        "preview_data": preview_data
    })


# 📤 Télécharger modèle Excel

def download_repository_template(request):
    # ✅ Colonnes attendues
    columns = [
        "indicator_code","country","spatial_dim","indicator_name",
        "dim1_type","dim1","dim2_type","dim2","dim3_type","dim3",
        "time_dim","alpha_value","numeric_value","publish_date"
    ]

    # ✅ Exemple de ligne
    data = [{
        "Indicator code": "MALARIA_INDIG",
        "Country": "Congo",
        "ISO": "CG",
        "Indicator name": "Malaria Indigenous Cases",
        "Dim1 type": "Age group",
        "Dim1": "15-49",
        "Dim2 type": "Gender",
        "Dim2": "Male",
        "Dim3 type": "Region",
        "Dim3": "Kinshasa",
        "Time_dim": "2026",
        "Alpha value": "A",
        "Numeric value": 1234,
        "Publish date": "2026-07-01"
    }]

    df = pd.DataFrame(data, columns=columns)

    # ✅ Export Excel
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="repository_template.xlsx"'
    df.to_excel(response, index=False)
    return response



# 📤 Exporter données en Excel
def export_repository_to_excel(request):
    response = HttpResponse(content_type="application/ms-excel")
    response["Content-Disposition"] = 'attachment; filename="repository_data.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "RepositoryData"

    headers = [
        "indicator_code","country","spatial_dim","indicator_name",
        "dim1_type","dim1","dim2_type","dim2","dim3_type","dim3",
        "time_dim","alpha_value","numeric_value","publish_date"
    ]
    ws.append(headers)

    for s in RepositoryIndicator.objects.all():
        ws.append([
            s.indicator_code, s.country, s.spatial_dim, s.indicator_name,
            s.dim1_type, s.dim1, s.dim2_type, s.dim2,
            s.dim3_type, s.dim3, s.time_dim,
            s.alpha_value, s.numeric_value, s.publish_date
        ])

    wb.save(response)
    return response



def import_excel_repository(request):
    REQUIRED_COLUMNS = [
        "indicator_code","country","spatial_dim","indicator_name",
        "dim1_type","dim1","dim2_type","dim2","dim3_type","dim3",
        "time_dim","alpha_value","numeric_value","publish_date"
    ]

    form = UploadExcelForm()
    preview_data_repository = None
    errors = []

    if request.method == "POST":
        form = UploadExcelForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES["file"]
            df = pd.read_excel(excel_file)

            # ✅ Vérifier colonnes obligatoires
            missing_cols = [col for col in REQUIRED_COLUMNS if col not in df.columns]
            if missing_cols:
                errors.append(f"Colonnes manquantes : {', '.join(missing_cols)}")

            # ✅ Conversion des dates pour éviter Timestamp
            if "publish_date" in df.columns:
                df["publish_date"] = pd.to_datetime(df["publish_date"], errors="coerce").dt.strftime("%Y-%m-%d")

            # ✅ Vérifier valeurs obligatoires
            for col in ["country", "indicator_code", "time_dim"]:
                if col in df.columns and df[col].isnull().any():
                    errors.append(f"Certaines lignes n’ont pas de valeur pour {col}")

            if not errors:
                preview_data_repository = df.head(20).to_dict(orient="records")

                # ✅ Conversion de toutes les colonnes datetime en string
                for col in df.select_dtypes(include=["datetime64[ns]"]).columns:
                    df[col] = df[col].dt.strftime("%Y-%m-%d")

                request.session["excel_data"] = df.to_dict(orient="records")

    return render(request, "mytvddata/pages/repository/import_excel.html", {
        "form": form,
        "preview_data_repository": preview_data_repository,
        "errors": errors,
    })




@csrf_exempt
def save_excel_repository(request):
    preview_data_repository = request.session.get("excel_data")
    report_data = []

    if preview_data_repository:
        created_count = 0
        updated_count = 0
        skipped_count = 0
        try:
            for c in preview_data_repository:
                # FK Country et Indicator
                country_obj, _ = Country.objects.get_or_create(
                    name=c.get("country"),
                    defaults={"population": c.get("population") or 0}
                    )
                
                indicator_obj = Indicator.objects.filter(indicator_name=c.get("indicator_name")).first()
                if not indicator_obj:
                    indicator_obj = Indicator.objects.create(indicator_name=c.get("indicator_name"))

              #  indicator_obj, _ = Indicator.objects.get_or_create(indicator_name=c.get("indicator_name"))

                # Conversion types
                time_dim_val = None
                if c.get("time_dim"):
                    try:
                        time_dim_val = float(c.get("time_dim"))
                    except:
                        time_dim_val = None

                publish_date_val = None
                if c.get("publish_date"):
                    try:
                        publish_date_val = datetime.datetime.strptime(c.get("publish_date"), "%Y-%m-%d").date()
                    except:
                        publish_date_val = None

                # ✅ Forcer une valeur si None
                if publish_date_val is None:
                   publish_date_val = datetime.today().date()

                # ✅ Vérification doublon sur toute la ligne
            #    existing = RepositoryIndicator.objects.filter(
                  #  country=country_obj,
                  #  indicator=indicator_obj,
                  #  indicator_code=c.get("indicator_code"),
                  #  spatial_dim=c.get("spatial_dim"),
                  #  dim1_type=c.get("dim1_type"),
                  #  dim1=c.get("dim1"),
                  #  dim2_type=c.get("dim2_type"),
                  #  dim2=c.get("dim2"),
                  #  dim3_type=c.get("dim3_type"),
                  #  dim3=c.get("dim3"),
                  #  time_dim=time_dim_val,
                  #  alpha_value=c.get("alpha_value"),
                  #  numeric_value=c.get("numeric_value"),
                  #  publish_date=publish_date_val,
               # ).first()

                existing = RepositoryIndicator.objects.filter(
                    country=country_obj,
                    spatial_dim=c.get("spatial_dim"),
                    indicator_code=c.get("indicator_code"),
                    time_dim=time_dim_val
                ).first()

                if existing:
                    # Mise à jour si différences
                    changed = False
                    for field in ["spatial_dim","dim1_type","dim1","dim2_type","dim2","dim3_type","dim3","alpha_value","numeric_value"]:
                        new_val = c.get(field)
                        if getattr(existing, field) != new_val:
                            setattr(existing, field, new_val)
                            changed = True
                    if changed:
                        existing.save()
                        updated_count += 1
                        report_data.append({"status": "Mis à jour", "raison": "Valeurs différentes", **c})
                    else:
                        skipped_count += 1
                        report_data.append({"status": "Ignoré", "raison": "Valeurs identiques", **c})
                else:
                    # Création si nouvelle ligne
                    RepositoryIndicator.objects.create(
                        country=country_obj,
                        indicator=indicator_obj,
                        indicator_code=c.get("indicator_code"),
                        spatial_dim=c.get("spatial_dim"),
                        dim1_type=c.get("dim1_type"),
                        dim1=c.get("dim1"),
                        dim2_type=c.get("dim2_type"),
                        dim2=c.get("dim2"),
                        dim3_type=c.get("dim3_type"),
                        dim3=c.get("dim3"),
                        time_dim=time_dim_val,
                        alpha_value=c.get("alpha_value"),
                        numeric_value=c.get("numeric_value"),
                        publish_date=publish_date_val,
                    )
                    created_count += 1
                    report_data.append({"status": "Créé", "raison": "Nouvelle ligne", **c})

            messages.success(request, f"✅ Import terminé : {created_count} créés, {updated_count} mis à jour, {skipped_count} doublons ignorés")
        except Exception as e:
            messages.error(request, f"Erreur lors de l’importation : {e}")
    else:
        messages.error(request, "Aucune donnée Excel à importer.")

    # ✅ Stocker le rapport complet dans la session
    request.session["import_report"] = report_data

    excel_data = request.session.get("excel_data")
    preview_data_repository = excel_data[:20] if excel_data else None
    
    #return redirect("datasets_repository")
    return render(request, "mytvddata/pages/repository/import_excel.html", {
    "form": UploadExcelForm(),
    "preview_data_repository": preview_data_repository,
    "created_count": created_count,
    "updated_count": updated_count,
    "skipped_count": skipped_count,
    "errors": [],
})




def export_import_report(request):
    # Récupérer les données du rapport depuis la session
    report_data = request.session.get("import_report")

    if not report_data:
        messages.error(request, "Aucun rapport disponible.")
        return redirect("import_excel_repository")

    # Créer un DataFrame
    df = pd.DataFrame(report_data)

    # Générer la réponse Excel
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="rapport_import.xlsx'
    df.to_excel(response, index=False)
    return response


# 📊 CRUD principal avec modals
def repository_add_data(request):
    form_select_repository = SelectRepositoryForm(request.GET or None)
    repositoryIndicators = RepositoryIndicator.objects.all().order_by("-publish_date")

  # 🔎 Filtres
    by_country = request.GET.get("country")
    by_indicator = request.GET.get("indicator")
    by_date = request.GET.get("date")
    by_subcomponent = request.GET.get("by_subcomponent")

    if by_country:
        repositoryIndicators = repositoryIndicators.filter(country__name__icontains=by_country)
    if by_indicator:
        repositoryIndicators = repositoryIndicators.filter(indicator_code__icontains=by_indicator)
    if by_date:
        repositoryIndicators = repositoryIndicators.filter(publish_date=by_date)
    if by_subcomponent:
        repositoryIndicators = repositoryIndicators.filter(indicator__subcomponent_indicator__icontains=by_subcomponent)

    form = RepositoryIndicatorForm()

       # ✅ Pagination
    paginator = Paginator(repositoryIndicators, 15)
    page_repos = paginator.get_page(request.GET.get('page'))

    try:
        repositoryIndicators = paginator.page(page_repos)
    except PageNotAnInteger:
        repositoryIndicators = paginator.page(1)
    except EmptyPage:
        repositoryIndicators = paginator.page(paginator.num_pages)



    # Charger les listes pour les selects
    countries = Country.objects.all().order_by("name")
    indicators = Indicator.objects.all().order_by("indicator_name")


    # Actions POST
    if request.method == "POST":
        if "save" in request.POST:
            pk = request.POST.get("save")
            if pk:
                instance = get_object_or_404(RepositoryIndicator, id=pk)
                form = RepositoryIndicatorForm(request.POST, instance=instance)
            else:
                form = RepositoryIndicatorForm(request.POST)

            if form.is_valid():
                form.save()
                messages.success(request, "✅ Donnée enregistrée ou mise à jour")
                return redirect("datasets_repository")

        elif "delete" in request.POST:
            pk = request.POST.get("delete")
            instance = get_object_or_404(RepositoryIndicator, id=pk)
            instance.delete()
            messages.success(request, "❌ Donnée supprimée")
            return redirect("datasets_repository")

        elif "edit" in request.POST:
            pk = request.POST.get("edit")
            instance = get_object_or_404(RepositoryIndicator, id=pk)
            form = RepositoryIndicatorForm(instance=instance)

    return render(request, "mytvddata/pages/repository/repository_dataset.html", {
        "repositoryIndicators": repositoryIndicators,
        "form": form,
        "page_repos":page_repos,
        "form_select_repository": form_select_repository,
        "countries": countries,
        "indicators": indicators,
    })


def repository_report(request):
    subComponents = Subcomponent.objects.all().order_by("component") 
    if request.method=='GET':
        st=request.GET.get('component')
        if st!=None:
            subComponents = Subcomponent.objects.filter(component__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(subComponents, num_of_items)
       
    try:
        subComponents = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        subComponents = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        subComponents = paginator.page(page)
            
    data={'subComponents': subComponents, 'paginator':paginator}
    
    return render(request, 'mytvddata/pages/repository/repository_report.html', data)

# Method for view single Survey DataSet page

def single_repository_page(request, pk):
    try:
        subComponent=Subcomponent.objects.get(id=pk) 
        subComponents =Subcomponent.objects.filter(Q(id__icontains=pk) & Q(subcomponent_indicator__indicator_repository__indicator_code__isnull=False)).values(
        'subcomponent_name',
        'subcomponent_indicator__indicator_repository__country__name',
        'subcomponent_indicator__indicator_repository__spatial_dim',
        'subcomponent_indicator__indicator_repository__indicator__indicator_name',
        'subcomponent_indicator__indicator_repository__indicator_code',
        'subcomponent_indicator__indicator_repository__dim1_type',
        'subcomponent_indicator__indicator_repository__dim1',
        'subcomponent_indicator__indicator_repository__dim2_type',
        'subcomponent_indicator__indicator_repository__dim2',
        'subcomponent_indicator__indicator_repository__dim3_type',
        'subcomponent_indicator__indicator_repository__dim3',
        'subcomponent_indicator__indicator_repository__time_dim',
        'subcomponent_indicator__indicator_repository__alpha_value',
        'subcomponent_indicator__indicator_repository__numeric_value',
        'subcomponent_indicator__indicator_repository__publish_date'
        
    ).order_by('-subcomponent_indicator__indicator_repository__publish_date')
        
        
        page = request.GET.get('page')
        num_of_items = 10
        # paginator of country 
        paginator = Paginator(subComponents, num_of_items)
        
        try:
            subComponents = paginator.page(page)
        except PageNotAnInteger:
            page = 1
            subComponents = paginator.page(page)
        except EmptyPage:
            page = paginator.num_pages
            subComponents = paginator.page(page)
        
        data={'subComponent':subComponent, 'subComponents':subComponents, 'paginator':paginator}   
    except:
        data={'message':'The repository DataSet you requested doesn t\' exist or isn t in the submitted results'}
    
    return render(request,'mytvddata/pages/repository/single_repository_page.html',data)

#Method to export REPOSITORY DATASET on Excel file
def export_repository_to_excel(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Saving_indicator_results.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Datawarehouse"

    # Add headers
    headers = ["sub_component","country", "spatial_dim", "indicator_name", "indicator_code", "dim1_type", "dim1", "dim2_type", "dim2", "dim3_type","dim3","time_dim","alpha_value","numeric_value","publish_date"]
    ws.append(headers)

    # Add data from the model
    repositoryIndicators = RepositoryIndicator.objects.all()
    for s in repositoryIndicators:
        ws.append([s.indicator.subcomponent.subcomponent_name,s.country.name, s.spatial_dim, s.indicator.indicator_name,s.indicator.indicator_code, s.dim1_type, s.dim1, 
                   s.dim2_type,  s.dim2,  s.dim3_type, s.dim3, s.time_dim, s.alpha_value, s.numeric_value, s.publish_date])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

 
 
# =========================
# EXPORT REPOSITORY PAGE (PRO)
# =========================
@login_required
def export_repository_page(request, pk):
 
    # ✅ Réponse HTTP (.xlsx moderne)
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=repository_data.xlsx'
 
    wb = Workbook()
    ws = wb.active
    ws.title = "Repository Data"
 
    # =========================
    # HEADERS
    # =========================
    headers = [
        "Subcomponent", "Country", "Spatial Dim",
        "Indicator Name", "Indicator Code",
        "Dim1 Type", "Dim1", "Dim2 Type", "Dim2",
        "Dim3 Type", "Dim3", "Time",
        "Alpha Value", "Numeric Value", "Publish Date"
    ]
 
    ws.append(headers)
 
    # =========================
    # QUERY OPTIMISÉE
    # =========================
    subcomponent = Subcomponent.objects.prefetch_related(
        'subcomponent_indicator__indicator_repository'
    ).get(pk=pk)
 
    # =========================
    # DATA
    # =========================
    for indicator in subcomponent.subcomponent_indicator.all():
        for repo_data in indicator.indicator_repository.all():
 
            ws.append([
                subcomponent.subcomponent_name,
                repo_data.country.name if repo_data.country else '',
                repo_data.spatial_dim,
                indicator.indicator_name,
                repo_data.indicator_code,
                repo_data.dim1_type,
                repo_data.dim1,
                repo_data.dim2_type,
                repo_data.dim2,
                repo_data.dim3_type,
                repo_data.dim3,
                repo_data.time_dim,
                repo_data.alpha_value,
                repo_data.numeric_value,
                repo_data.publish_date
            ])
 
    # =========================
    # AUTO WIDTH (UX)
    # =========================
    for col in ws.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
 
        for cell in col:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
 
        ws.column_dimensions[col_letter].width = max_length + 2
 
    # ✅ SAVE
    wb.save(response)
 
    return response
 


## METHOD EXPORT THE FILTERING DATA UNIT IN EXCEL FORMAT
def export_to_excel_repositoring(request,by_subcomponent,end_day):
    #startdate
    """
    Une vue Django qui prend l'objet request et trois autres paramètres.
    """
 #Q(subcomponent_indicator__indicator_repository__publish_date__gte=start_day)&
       
    sub_component=Subcomponent.objects.get(id=by_subcomponent)
    data_component = Subcomponent.objects.all().filter(Q(id__icontains=by_subcomponent)&  Q(subcomponent_indicator__indicator_repository__publish_date__lte=end_day)).values(
        'subcomponent_name',
        'subcomponent_indicator__indicator_repository__country__name',
        'subcomponent_indicator__indicator_repository__spatial_dim',
        'subcomponent_indicator__indicator_repository__indicator__indicator_name',
        'subcomponent_indicator__indicator_repository__indicator_code',
        'subcomponent_indicator__indicator_repository__dim1_type',
        'subcomponent_indicator__indicator_repository__dim1',
        'subcomponent_indicator__indicator_repository__dim2_type',
        'subcomponent_indicator__indicator_repository__dim2',
        'subcomponent_indicator__indicator_repository__dim3_type',
        'subcomponent_indicator__indicator_repository__dim3',
        'subcomponent_indicator__indicator_repository__time_dim',
        'subcomponent_indicator__indicator_repository__alpha_value',
        'subcomponent_indicator__indicator_repository__numeric_value',
        'subcomponent_indicator__indicator_repository__publish_date'
        ).order_by('-subcomponent_indicator__indicator_repository__publish_date')
    
   # number_register =data_survey.annotate(total_sample=Count('project_surveyData__level_1')).values('total_sample')
    # Create fill
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="'+'Saving_dataWarehouse_'+sub_component.subcomponent_name+"_"+end_day+'.xlsx"'  
    workbook=Workbook()
    worksheet = workbook.active
    
    # Edit page Setup
    worksheet.page_setup.orientation = worksheet.ORIENTATION_LANDSCAPE
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_TABLOID
    worksheet.page_setup.fitToHeight = 0
    worksheet.page_setup.fitToWidth = 1
    
    worksheet.merge_cells('A1:F1')
    worksheet.merge_cells('A2:F2')
    worksheet.freeze_panes = "F4"
    worksheet.auto_filter.ref = "A3:J5000"
    worksheet["H1"] = "TOTAL REGISTER"
    worksheet["I1"] = '=COUNTA(C4:C5000)'
    worksheet["I1"].font = Font(bold=True,color="0000ff", size=14)
    worksheet["I1"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["F2"] = "NEW RISK"
 #   worksheet["G2"] = '=COUNTIF(G4:G200, "New")'
 #   worksheet["G2"].font = Font(bold=True, color="00FF0000", size=14, italic=True)
 #   worksheet["G2"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["H2"] = "EXISTING RISK"
 #   worksheet["I2"] = '=COUNTIF(G4:G200, "Existing")'
 #   worksheet["I2"].font = Font(bold=True, color="00993300", size=14, italic=True)
 #   worksheet["I2"].alignment = Alignment(horizontal ="center", vertical="center")
 #   worksheet["J2"] = "CLOSED"
 #   worksheet["K2"] = '=COUNTIF(G4:G200, "closed")'
 #   worksheet["K2"].font = Font(bold=True, color="ff00ff", size=14, italic=True)
 #   worksheet["K2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #   worksheet["L2"] = "MODERATE"
 #   worksheet["M2"] = '=COUNTIF(J4:J200, "Moderate")'
 #   worksheet["M2"].font = Font(bold=True, color="0000CCFF", size=14, italic=True)
 #   worksheet["M2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #   worksheet["N2"] = "CRITICAL"
 #   worksheet["O2"] = '=COUNTIF(J4:J200, "Critical")'
 #   worksheet["O2"].font = Font(bold=True, color="00993366", size=14, italic=True)
 #   worksheet["O2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #   worksheet["P2"] = "VERY CRITICAL"
 ##   worksheet["Q2"] = '=COUNTIF(J4:J200, "Very_critical")'
 #   worksheet["Q2"].font = Font(bold=True, color="00FF6600", size=14, italic=True)
 #   worksheet["Q2"].alignment = Alignment(horizontal ="center", vertical="center")
    
 #  worksheet["R1"] = "MONITORING (%)"
 #   worksheet["S1"] ='=AVERAGE(R4:R20)'
 #   worksheet["S1"].font = Font(bold=True, color="0000FF", size=14, italic=True)
 #   worksheet["S1"].alignment = Alignment(horizontal ="center", vertical="center")
    #worksheet.cell['S1'].number_format = '0.00%'
 #   worksheet.conditional_formatting.add("S1", color_scale_rule)
    
   
  
  # Again, let's add this gradient to the star ratings, column "H"
  #  worksheet.conditional_formatting.add("R4:R20", color_scale_rule)
   
    first_cell = worksheet['A1']
    first_cell.value ="Repository: " +""+ sub_component.component.component_name +" and End date: "+end_day
    first_cell.fill = PatternFill("solid",fgColor="246ba1")
    first_cell.font =  Font(bold=True, color="F7F6FA", size=16)
    first_cell.alignment = Alignment(horizontal ="center", vertical="center")
     
    second_cell = worksheet['A2']
    second_cell.value = sub_component.subcomponent_name
    second_cell.font = Font(bold=True, color="246ba1", size=16)
    second_cell.alignment = Alignment(horizontal ="center", vertical="center")
    
   
        
    worksheet.title='sub_component.component.component_name' + "" +"_"+end_day
     # Define the titles for columns
    columns = [
               "sub_component","country", "spatial_dim", "indicator_name",
               "indicator_code", "dim1_type", "dim1", "dim2_type", "dim2", 
               "dim3_type","dim3","time_dim","alpha_value","numeric_value",
               "publish_date"
                ]
 
    row_num = 3
    
    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value= column_title
        cell.fill = PatternFill("solid", fgColor="50c878")
        cell.font = Font(bold=True, color="F7F6FA")
        cell.alignment = Alignment(horizontal ="center", vertical="center",wrap_text=False)
        third_cell=worksheet['F3']
        third_cell.alignment =Alignment(horizontal="center", vertical="center",wrap_text=False)
        
    for subComponent in data_component:
        row_num +=1    
        #Define the data for each cell in the row
        row=[subComponent['subcomponent_name'],
             subComponent['subcomponent_indicator__indicator_repository__country__name'],
             subComponent['subcomponent_indicator__indicator_repository__spatial_dim'],
             subComponent['subcomponent_indicator__indicator_repository__indicator__indicator_name'],
             subComponent['subcomponent_indicator__indicator_repository__indicator_code'],
             subComponent['subcomponent_indicator__indicator_repository__dim1_type'],
             subComponent['subcomponent_indicator__indicator_repository__dim1'],
             subComponent['subcomponent_indicator__indicator_repository__dim2_type'],
             subComponent['subcomponent_indicator__indicator_repository__dim2'],
             subComponent['subcomponent_indicator__indicator_repository__dim3_type'],
             subComponent['subcomponent_indicator__indicator_repository__dim3'],
             subComponent['subcomponent_indicator__indicator_repository__time_dim'],
             subComponent['subcomponent_indicator__indicator_repository__alpha_value'],
             subComponent['subcomponent_indicator__indicator_repository__numeric_value'],
             subComponent['subcomponent_indicator__indicator_repository__publish_date']
             ]
                   
          #Assign the data for each cell in the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment =Alignment(vertical="center",wrap_text=False)
    #          if isinstance(cell_value, decimal.Decimal):
    #              cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
    workbook.save(response)
    return response

    
def upload_success_repo(request):
    return render(request, 'mytvddata/pages/repository/upload_success_repo.html')


def dataset_list_reposi(request):
    repository_indicators = RepositoryIndicator.objects.all()
    return render(request, 'mytvddata/pages/repository/dataset_list_reposi.html', {'repository_indicators': repository_indicators})



#### API GHO IMPORT DATA (GOOD)--------------------------------------------------

@csrf_exempt
def import_api_gho_data(api_url, selected_fields):
    response = requests.get(api_url, timeout=15)
    response.raise_for_status()
    details = response.json().get("value", [])

    if not details:
        return "Aucune donnée trouvée."

    # Mapping API → DB
    mappings = {
        m.api_field: m.db_field
        for m in ApiFieldMapping.objects.filter(api_field__in=selected_fields)
    }

    inserted, updated = 0, 0
    for item in details:
        row = {}
        for api_field in selected_fields:
            db_field = mappings.get(api_field)
            if db_field:
                value = item.get(api_field)
                if "date" in db_field.lower() and value:
                    try:
                        value = datetime.strptime(value[:10], "%Y-%m-%d").date()
                    except Exception:
                        pass
                row[db_field] = value

        if row:
            obj, created = StoreAPI.objects.update_or_create(
                api_id=item.get("Id"),  # clé unique
                defaults=row
            )
            if created:
                inserted += 1
            else:
                updated += 1

    return f"{inserted} nouveaux enregistrements, {updated} mis à jour."


@csrf_exempt
def select_api_gho_fields(request):
    # Champs autorisés
    ALLOWED_FIELDS = [
        "Id", "IndicatorCode", "SpatialDim", "Dim1Type", "Dim1",
        "TimeDim", "Dim2Type", "Dim2", "Dim3Type", "Dim3",
        "Value", "NumericValue", "Date"
    ]
    available_fields, sample_data = [], []
    db_fields = [f.name for f in StoreAPI._meta.get_fields() if f.concrete and not f.auto_created]

    url_form = SelectUrlForm(request.POST or None)
    action = request.POST.get("action")

    if request.method == "POST":
        # Étape 1 : Détection
        if action == "validate" and url_form.is_valid():
            indicator = url_form.cleaned_data["indicator_code"]
            start, end = url_form.cleaned_data["start"], url_form.cleaned_data["end"]

            api_url = f"https://ghoapi.azureedge.net/api/{indicator}" \
                      f"?$filter=ParentLocationCode eq 'AFR' " \
                      f"and date(TimeDimensionBegin) ge {start} " \
                      f"and date(TimeDimensionBegin) lt {end}"

            response = requests.get(api_url, timeout=10)
            details = response.json().get("value", [])

            if details:
                first_item = details[0]
                available_fields = [{"id": f, "field_name": f} for f in ALLOWED_FIELDS if f in first_item]
                sample_data = [{f: row.get(f) for f in ALLOWED_FIELDS if f in row} for row in details[:3]]

                request.session["api_url"] = api_url
                request.session["available_fields"] = [f["field_name"] for f in available_fields]

        # Étape 2 : Sélection
        elif action == "filter":
            selected_fields = request.POST.getlist("selected_fields")
            request.session["available_fields"] = selected_fields

            api_url = request.session.get("api_url")
            if api_url:
                response = requests.get(api_url, timeout=10)
                details = response.json().get("value", [])
                sample_data = [{f: row.get(f) for f in selected_fields} for row in details[:3]]

            available_fields = [{"id": f, "field_name": f} for f in selected_fields]
            messages.success(request, f"{len(selected_fields)} champ(s) sélectionné(s).")

        # Étape 3 : Mapping + Import
        elif action == "import":
            api_url = request.session.get("api_url")
            selected_fields = request.session.get("available_fields", [])

            if not api_url or not selected_fields:
                messages.error(request, "⚠️ Aucun champ sélectionné ou URL API manquante.")
            else:
                for field in selected_fields:
                    mapping_value = request.POST.get(f"mapping_{field}")
                    if mapping_value:
                        ApiFieldMapping.objects.update_or_create(
                            api_field=field,
                            defaults={"db_field": mapping_value}
                        )
                message = import_api_gho_data(api_url, selected_fields)
                messages.success(request, message)

        # Étape 4 : Reset
        elif action == "reset_mapping":
            ApiFieldMapping.objects.all().delete()
            request.session.pop("available_fields", None)
            messages.info(request, "Le mapping a été réinitialisé.")

    else:
        # Rechargement depuis la session
        if "available_fields" in request.session:
            available_fields = [{"id": f, "field_name": f} for f in request.session["available_fields"]]
        if "api_url" in request.session:
            api_url = request.session["api_url"]
            try:
                response = requests.get(api_url, timeout=10)
                sample_data = response.json().get("value", [])[:3]
            except Exception:
                sample_data = []

          # Vérifier si la base contient des données
    records_count = StoreAPI.objects.count()

    return render(request, "mytvddata/cover/select_fields_gho.html", {
        "form": url_form,
        "available_fields": available_fields,
        "sample_data": sample_data,
        "db_fields": db_fields,
        "records_exist": records_count > 0,   # booléen
         "records_count": records_count        # nombre total
    })


@csrf_exempt
def view_imported_data_gho(request):
    records = StoreAPI.objects.all()
    return render(request, "mytvddata/cover/imported_data.html", {"records": records})


@csrf_exempt
def view_api_field_mapping(request):
    mappings = ApiFieldMapping.objects.all().order_by("api_field")
    return render(request, "mytvddata/cover/gho_mapping.html", {"mappings": mappings})


@csrf_exempt
def load_api_gho_other(request):
    api_url = request.session.get("api_url")  # récupère l’URL saisie par l’utilisateur
    if not api_url:
        messages.error(request, "Veuillez d’abord saisir une URL API.")
        return redirect("select_api_fields")

    headers = {"Content-Type": "application/json"}
    allowed_fields = list(ApiFieldConfig.objects.filter(enabled=True).values_list("field_name", flat=True))

    try:
        response = requests.get(api_url, headers=headers, verify=False)
        response.raise_for_status()
        payload = response.json().get("value", [])

        for item in payload:
            filtered_data = {f: item.get(f) for f in allowed_fields}

            StoreAPI.objects.update_or_create(
                api_id=item.get("Id"),
                defaults={
                    "indicator_code": filtered_data.get("IndicatorCode"),
                    "country_code": filtered_data.get("SpatialDim"),
                    "parent_location": filtered_data.get("ParentLocation"),
                    "time_dim": filtered_data.get("TimeDim"),
                    "value": filtered_data.get("Value"),
                    "numeric_value": filtered_data.get("NumericValue"),
                    "date": filtered_data.get("Date"),
                }
            )
        messages.success(request, "API data saved successfully!")

    except Exception as e:
        messages.error(request, f"Error fetching API: {e}")

    data_from_db = StoreAPI.objects.all().order_by("-date")
    return render(request, "mytvddata/cover/api_gho_data_other.html", {"api_data": data_from_db})


@csrf_exempt
def view_api_field_mapping(request):
    mappings = ApiFieldMapping.objects.all().order_by("api_field")
    return render(request, "mytvddata/cover/gho_mapping.html", {
        "mappings": mappings
    })


######################################################











# Affiche la table des données API GHO et gère les actions CRUD



def display_api_data(request):
    form_select_api = SelectAPIForm(request.GET or None)
    by_indicator = request.GET.get('by_indicator')
    start_day = request.GET.get('startdate')
    end_day = request.GET.get('end')
    indicator_code = request.GET.get('indicator_code')

    # ✅ Base queryset
    storeAPIs = StoreAPI.objects.all().order_by("-publish_date")

    # ✅ Filtrage
    if indicator_code:
        storeAPIs = storeAPIs.filter(indicator_code__icontains=indicator_code)

    if by_indicator and start_day and end_day:
        storeAPIs = storeAPIs.filter(
            Q(indicator_code__icontains=by_indicator),
            Q(publish_date__gte=start_day),
            Q(publish_date__lte=end_day)
        )

    # ✅ Pagination
    paginator = Paginator(storeAPIs, 15)
    page_api = paginator.get_page(request.GET.get('page'))

    try:
        storeAPIs = paginator.page(page_api)
    except PageNotAnInteger:
        storeAPIs = paginator.page(1)
    except EmptyPage:
        storeAPIs = paginator.page(paginator.num_pages)

    # ✅ Formulaire par défaut
    form = StoreAPIForm()

    # ✅ Actions POST
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if pk:
                instance = get_object_or_404(StoreAPI, id=pk)
                form = StoreAPIForm(request.POST, instance=instance)
            else:
                form = StoreAPIForm(request.POST)

            if form.is_valid():
                form.save()
                messages.success(request, "✅ API data saved or updated successfully")
                return redirect("display_api_data")

        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            instance = get_object_or_404(StoreAPI, id=pk)
            instance.delete()
            messages.success(request, "❌ API data deleted successfully")
            return redirect("display_api_data")

        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            instance = get_object_or_404(StoreAPI, id=pk)
            form = StoreAPIForm(instance=instance)

    context = {
        "storeAPIs": storeAPIs,
        "page_api": page_api,
        "form": form,
        "form_select_api": form_select_api,
        "by_indicator": by_indicator,
        "start_day": start_day,
        "end_day": end_day,
        "title": "home",
    }
    # data_from_db = StoreAPI.objects.all().order_by('-created_at')
    return render(request, 'mytvddata/epidata/api_dataset.html', context)



@csrf_exempt
def load_api_json(request):
    available_fields = []
    sample_data = []
    db_fields = [f.name for f in StoreAPI._meta.get_fields() if f.concrete and not f.auto_created]

    form = SelectUrlForm(request.POST or None)

    if request.method == "POST" and form.is_valid():
     
        api_url = request.session.get("api_url")
        if not api_url:
            messages.error(request, "Veuillez d’abord détecter les champs API avant de charger les données.")
            return redirect("select_api_fields")





        action = request.POST.get("action")

        try:
            response = requests.get(api_url, timeout=10)
            response.raise_for_status()
            payload = response.json()
            details = payload.get("value", [])

            if details:
                # Champs disponibles
                first_item = details[0]
                available_fields = [{"id": key, "field_name": key} for key in first_item.keys()]

                # Aperçu des 3 premières lignes
                sample_data = details[:3]

                # Sauvegarde du mapping choisi
                for field in available_fields:
                    mapping_value = request.POST.get(f"mapping_{field['id']}")
                    if mapping_value:
                        ApiFieldMapping.objects.update_or_create(
                            api_field=field["field_name"],
                            defaults={"db_field": mapping_value}
                        )

                # Import si demandé
                if action == "import":
                    message = import_api_data(api_url)
                    messages.success(request, f"{message} Cliquez ci-dessous pour voir les données importées.")

        except Exception as e:
            messages.error(request, f"Erreur lors de la récupération : {e}")

    context = {
        "form": form,
        "available_fields": available_fields,
        "sample_data": sample_data,
        "db_fields": db_fields,
    }
    return render(request, "mytvddata/cover/api_data_list.html", context)

@csrf_exempt
def import_api_data(api_url):
    response = requests.get(api_url, timeout=10)
    response.raise_for_status()
    payload = response.json()
    details = payload.get("value", [])

    if not details:
        return "Aucune donnée trouvée."

    mappings = {m.api_field: m.db_field for m in ApiFieldMapping.objects.all()}

    for item in details:
        record_data = {}
        for api_field, db_field in mappings.items():
            if db_field and api_field in item:
                record_data[db_field] = item[api_field]

        StoreAPI.objects.create(**record_data)

    return f"{len(details)} lignes importées avec succès."


 
@csrf_exempt
def loads_api(request): # GOOD Method to store data from API to MySQL database USE WITH POSTMAN PLATFORM

    payload=json.loads(request.body)
    data = payload.get('value')
        
        #    url = "https://api.example.com/users"  # Replace with your API
        #    response = requests.get(url)
    if request.method == 'POST':
            
     #   if data.status_code == 200:
     #       data = response.json()  # JSON → Python dict/list
            
            for item in data:
                # Insert into MySQL via Django ORM
                ApiData.objects.update_or_create(
                    api_id=item['Id'],  # unique field
                    defaults={
                        "IndicatorCode": item["IndicatorCode"],
                        "SpatialDimType": item["SpatialDimType"],
                        "SpatialDim": item["SpatialDim"],
                        "TimeDimType": item["TimeDimType"],
                        "ParentLocationCode": item["ParentLocationCode"],
                        "ParentLocation": item["ParentLocation"],
                        "Dim1Type": item["Dim1Type"],
                        "Dim1": item["Dim1"],
                        "TimeDim": item["TimeDim"],
                        "Dim2Type": item["Dim2Type"],
                        "Dim2": item["Dim2"],
                        "Dim3Type": item["Dim3Type"],
                        "Dim3": item["Dim3"],
                        "DataSourceDim": item["DataSourceDim"],
                        "Value": item["Value"],
                        "NumericValue": item["NumericValue"],
                        "Low": item["Low"],
                        "High": item["High"],
                        "Comments": item["Comments"],
                        "Date": item["Date"],
                        "TimeDimensionValue": item["TimeDimensionValue"],
                        "TimeDimensionBegin": item["TimeDimensionBegin"],
                        "TimeDimensionEnd": item["TimeDimensionEnd"]
                    }
                )
    else:
        print("Failed to fetch data:")
        
    return HttpResponse("Data successfully fetched and stored!")    


# Optional: View to display the fetched GHO from your database
def fetch_api_stored(request):
    response={}  
    payload={}
    header={"Content-Type":"applica/json",
            "Accept-Encoding":"Deflate"}
    
    if request.method=='GET':
        st=request.GET.get("api_url")  # 1. Appeler l'API JSON
        if st!=None:
            response = requests.get(st,header, verify=False)
            response.raise_for_status()  # Lève une exception pour les codes d'état d'erreur

            if response.status_code==200:  # Vérifier que la requête est OK
               payload=json.loads(response.text)
             #  details = payload.get('value')
               return JsonResponse({
                     'payload':payload
               }, safe=False)
                
            else:
                messages.info(request, "Erreur lors de l'appel AP:", response.status_code)
    return JsonResponse({
                     'response':response
               }, safe=False)
                
  #  return HttpResponse("Données de l'API chargées avec succès dans la base de données!")
 
        
# Method for upload pdf document or report in pdf format
#@lockdown()
#@permission_required('myuhp.add_reportsave',raise_exception=True)
def report_upload(request):
    img_report_obj = None
    img_report_obj_all = ReportSave.objects.all()

    if request.method == "POST":
        form = ReportSaveForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, "Report uploaded successfully")
            img_report_obj = form.instance
        else:
            messages.error(request, "Upload failed. Please check the file format.")
    else:
        form = ReportSaveForm()

    return render(request, "mytvddata/pages/report/report_upload.html", {
        "form": form,
        "img_report_obj": img_report_obj,
        "img_report_obj_all": img_report_obj_all,
    })



# Method for views all reports upload in the system
def views_report(request):
    reportSaves = ReportSave.objects.all().order_by("title_rep") 
    if request.method=='GET':
        st=request.GET.get('title_rep')
        if st!=None:
            reportSaves= ReportSave.objects.filter(title_rep__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(reportSaves, num_of_items)
       
    try:
        reportSaves = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        reportSaves = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        reportSaves = paginator.page(page)
            
    data={'reportSaves': reportSaves, 'paginator':paginator}
    
    return render(request, 'mytvddata/pages/report/report_list.html', data)

# Method to put  into the table all reports upload in the system
#@login_required
#@permission_required('myuhp.view_reportsave',raise_exception=True)
def index_report(request):
    reportSaves = ReportSave.objects.all().order_by("title_rep") 
    if request.method=='GET':
        st=request.GET.get('title_rep')
        if st!=None:
            reportSaves= ReportSave.objects.filter(title_rep__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 10
    # paginator of country 
    paginator = Paginator(reportSaves, num_of_items)
       
    try:
        reportSaves = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        reportSaves = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        reportSaves = paginator.page(page)
            
    data={'reportSaves': reportSaves, 'paginator':paginator}
    
    return render(request, 'mytvddata/pages/report/report_index.html', data)

# Edit report thad added in the system
#@login_required
#@permission_required('myuhp.add_reportsave',raise_exception=True)



def edit_report(request, pk):
    reportSave = get_object_or_404(ReportSave, pk=pk)
    if request.method == "POST":
        form = ReportSaveForm(request.POST, request.FILES, instance=reportSave)
        if form.is_valid():
            form.save()
            messages.success(request, "Report updated successfully")
    else:
        form = ReportSaveForm(instance=reportSave)

    # ✅ Construire l’URL absolue côté Python
    file_url = request.build_absolute_uri(reportSave.file_rep.url) if reportSave.file_rep else None

    return render(request, "mytvddata/pages/report/report_edit.html", {
        "form": form,
        "reportSave": reportSave,
        "file_url": file_url,   # on passe l’URL absolue au template
    })

# Method for upload document or image
#@lockdown()
#@permission_required('myuhp.add_docsave',raise_exception=True)
def docSave_upload(request):
    form =None
    if request.method == "POST":
        form = DocSaveForm(request.POST, request.FILES)
       # print(form)
        if form.is_valid():
            form.save()
            # Get the current instance object to display in the template
            img_obj = form.instance
        return render(request, 'mytvddata/pages/report/doc_upload.html', {'form': form, 'img_obj': img_obj})
    else:
        form = DocSaveForm()
        img_obj_all = DocSave.objects.all()
     
    return render(request,'mytvddata/pages/report/doc_upload.html', {'form': form, 'img_obj_all':img_obj_all})


## METHOD EXPORT THE FILTERING DATA UNIT IN EXCEL FORMAT
def export_to_excel_API(request,by_indicator,end_day):
   
 #Q(subcomponent_indicator__indicator_repository__publish_date__gte=start_day)&
       
   # sub_component=StoreAPI.objects.get(indicator_code=by_indicator)
    data_component = StoreAPI.objects.all().filter(Q(indicator_code__icontains=by_indicator)&  Q(publish_date__lte=end_day)).values(
        'api_id',
        'country_code',
        'indicator_code',
        'dim1_type',
        'dim1',
        'dim2_type',
        'dim2',
        'dim3_type',
        'dim3',
        'time_dim',
        'alpha_value',
        'numeric_value',
        'publish_date'
        ).order_by('-publish_date')
    
   # number_register =data_survey.annotate(total_sample=Count('project_surveyData__level_1')).values('total_sample')
    # Create fill
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="'+'Saving_api_'+by_indicator+"_"+end_day+'.xlsx"'  
    workbook=Workbook()
    worksheet = workbook.active
    
    # Edit page Setup
    worksheet.page_setup.orientation = worksheet.ORIENTATION_LANDSCAPE
    worksheet.page_setup.paperSize = worksheet.PAPERSIZE_TABLOID
    worksheet.page_setup.fitToHeight = 0
    worksheet.page_setup.fitToWidth = 1
    
    worksheet.merge_cells('A1:F1')
    worksheet.merge_cells('A2:F2')
    worksheet.freeze_panes = "F4"
    worksheet.auto_filter.ref = "A3:J5000"
    worksheet["H1"] = "TOTAL REGISTER"
    worksheet["I1"] = '=COUNTA(C4:C5000)'
    worksheet["I1"].font = Font(bold=True,color="0000ff", size=14)
    worksheet["I1"].alignment = Alignment(horizontal ="center", vertical="center")
 
   
    first_cell = worksheet['A1']
    first_cell.value ="API: " +""+ by_indicator +" and End date: "+end_day
    first_cell.fill = PatternFill("solid",fgColor="246ba1")
    first_cell.font =  Font(bold=True, color="F7F6FA", size=16)
    first_cell.alignment = Alignment(horizontal ="center", vertical="center")
     
   
        
    worksheet.title='by_indicator' + "" +"_"+end_day
     # Define the titles for columns
    columns = [
               "api_id", "country_code", "indicator_code","dim1_type",
                "dim1","dim2_type","dim2","dim3_type","dim3","time_dim",
                "alpha_value","numeric_value","publish_date"
                ]
 
    row_num = 3
    
    # Assign the titles for each cell of the header
    for col_num, column_title in enumerate(columns, 1):
        cell = worksheet.cell(row=row_num, column=col_num)
        cell.value= column_title
        cell.fill = PatternFill("solid", fgColor="50c878")
        cell.font = Font(bold=True, color="F7F6FA")
        cell.alignment = Alignment(horizontal ="center", vertical="center",wrap_text=False)
        third_cell=worksheet['F3']
        third_cell.alignment =Alignment(horizontal="center", vertical="center",wrap_text=False)
        
    for subComponent in data_component:
        row_num +=1    
        #Define the data for each cell in the row
        row=[subComponent['api_id'],
             subComponent['country_code'],
             subComponent['indicator_code'],
             subComponent['dim1_type'],
             subComponent['dim1'],
             subComponent['dim2_type'],
             subComponent['dim2'],
             subComponent['dim3_type'],
             subComponent['dim3'],
             subComponent['time_dim'],
             subComponent['alpha_value'],
             subComponent['numeric_value'],
             subComponent['publish_date']
             ]
                   
          #Assign the data for each cell in the row
        for col_num, cell_value in enumerate(row, 1):
            cell = worksheet.cell(row=row_num, column=col_num)
            cell.value = cell_value
            cell.alignment =Alignment(vertical="center",wrap_text=False)
    #          if isinstance(cell_value, decimal.Decimal):
    #              cell.number_format = numbers.FORMAT_NUMBER_COMMA_SEPARATED1
    workbook.save(response)
    return response


# def component(request):
#    form=ComponentForm()
#    context={'form':form}
#    return render(request, 'mytvddata/pages/component/component.html', context)
    
#def subcomponents(request):
#    form=ComponentForm(request.GET)
#    return HttpResponse(form['subcomponents'])
    
# Method for CRUD Component


def component_add(request):
    context = {}
    form = ComponentsForm()
    components = Component.objects.all()

    if request.method == 'GET':
        st = request.GET.get('component_name')
        if st:
            components = Component.objects.filter(component_name__icontains=st)

    page = request.GET.get('page')
    num_of_items = 10
    paginator = Paginator(components, num_of_items)

    try:
        components = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        components = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        components = paginator.page(page)

    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form = ComponentsForm(request.POST)
            else:
                component = Component.objects.get(id=pk)
                form = ComponentsForm(request.POST, instance=component)
            if form.is_valid():
                form.save()
                messages.success(request, 'Component saved or updated successfully')
            return redirect('component_add')  # 🔑 redirection après POST

        elif 'delete' in request.POST:
            pk = request.POST.get('delete')
            component = Component.objects.get(id=pk)
            component.delete()
            messages.success(request, 'Component deleted successfully')
            return redirect('component_add')  # 🔑 redirection après POST

        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            component = Component.objects.get(id=pk)
            form = ComponentsForm(instance=component)

    context['components'] = components
    context['paginator'] = paginator
    context['title'] = 'home'
    context['form'] = form

    return render(request, 'mytvddata/pages/component/component.html', context)


#Method to export ALL COMPONENT on Excel file
def export_component(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="Component.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "Component"

    # Add headers
    headers = ["id","component_name"]
    ws.append(headers)

    # Add data from the model
    components = Component.objects.all()
    for component in components:
        ws.append([component.pk, component.component_name])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response


# Method for CRUD Component
def subcomponent_add(request):
    context={}
    form = SubcomponentsForm()
    subcomponents = Subcomponent.objects.all()
    
    if request.method=='GET':
        st=request.GET.get('subcomponent_name')
        if st!=None:
            subcomponents= Subcomponent.objects.filter(subcomponent_name__icontains=st)
    
    page = request.GET.get('page')
    num_of_items = 100
    # paginator of country 
    paginator = Paginator(subcomponents, num_of_items)
       
    try:
        subcomponents = paginator.page(page)
    except PageNotAnInteger:
        page = 1
        subcomponents = paginator.page(page)
    except EmptyPage:
        page = paginator.num_pages
        subcomponents = paginator.page(page)
            
    context['subcomponents'] = subcomponents
  
    context['paginator']=paginator
    context['title'] ='home'
    if request.method == 'POST':
        if 'save' in request.POST:
            pk = request.POST.get('save')
            if not pk:
                form=SubcomponentsForm(request.POST)
            else:
                subcomponent = Subcomponent.objects.get(id = pk)
                form = SubcomponentsForm(request.POST ,instance = subcomponent) 
            if form.is_valid():
                form.save()
                messages.success(request, 'Disease saved or updated successfully')
            return redirect('subcomponent_add')  # 🔑 redirection après POST

        elif 'delete' in request.POST:
         #   if request.user.groups.filter(name='Manager and delete').exists(): ## Condition security
                pk = request.POST.get('delete')
                subcomponent = Subcomponent.objects.get(id = pk)
                subcomponent.delete()
                messages.success(request, 'Sub Component deleted successfully')
                return redirect('subcomponent_add')  # 🔑 redirection après POST
        #    else:
        #        return render(request, 'pages/examples/403.html')
    
        elif 'edit' in request.POST:
            pk = request.POST.get('edit')
            subcomponent = Subcomponent.objects.get(id = pk)
            form = SubcomponentsForm(instance = subcomponent)
          
    context['form'] = form
    context['paginator'] = paginator
    context['title'] = 'home'
  
    return render(request, 'mytvddata/pages/component/subcomponent.html', context)


#Method to export ALL SUB-COMPONENT on Excel file
def export_subcomponent(request):
    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = 'attachment; filename="SubComponent.xlsx"'

    wb = Workbook()
    ws = wb.active
    ws.title = "SubComponent"

    # Add headers
    headers = ["id","component","subcomponent_name"]
    ws.append(headers)

    # Add data from the model
    subcomponents = Subcomponent.objects.all()
    for subcomponent in subcomponents:
        ws.append([subcomponent.pk, subcomponent.component.component_name,subcomponent.subcomponent_name])

    # Save the workbook to the HttpResponse
    wb.save(response)
    return response

# CREATE THE JSON DATA FOR INDICATOR
@csrf_exempt
def json_indicator(request):
    data=list(Indicator.objects.values('indicator_code','indicator_name','indicator_target','indicator_metric','type_indicator','ref_data',
                                       'indicator_unit','indicator_source','category_indicator','forecasting_indicator','performance_indicator','subcomponent__subcomponent_name','subcomponent__component__component_name'))
    return JsonResponse(data, safe=False)


# CREATE THE JSON DATA FOR STOREAPI
@csrf_exempt
def json_storeAPI(request):
    data_storeAPI=list(StoreAPI.objects.values('indicator_code','country_code','time_dim','dim1_type','dim1',
                                               'dim2_type','dim2','dim3_type','dim3','alpha_value','numeric_value'))
    return JsonResponse(data_storeAPI, safe=False)


# CREATE THE JSON DATA FOR REPOSITORY
@csrf_exempt
def json_repository(request):
    data_repository=list(RepositoryIndicator.objects.values('indicator__subcomponent__component__component_name','indicator__subcomponent__subcomponent_name','country__name','spatial_dim', 'indicator__indicator_name',
                                                            'indicator__indicator_code','dim1_type', 'dim1', 'dim2_type', 'dim2',  'dim3_type', 'dim3', 'time_dim', 
                                                            'alpha_value', 'numeric_value', 'publish_date'))  
    return JsonResponse(data_repository, safe=False)


# CREATE THE JSON DATA FOR COUNTRY
@csrf_exempt
def json_country(request):
    data_country=list(Country.objects.values('flag','cca3','name','official','capital','subregion',
                                             'area','population','languages','ref_data','data_source','country_class'))
    return JsonResponse(data_country, safe=False)


# CREATE THE JSON DATA FOR LOCATION GPS
@csrf_exempt
def json_location(request):
    data_location=list(LocationCountry.objects.values('iso3','name','latitude','longitude'))
    return JsonResponse(data_location, safe=False)

# CREATE THE JSON DATA FOR COMPONENT
@csrf_exempt
def json_component(request):
    data_component=list(Component.objects.values('id','component_name'))
    return JsonResponse(data_component, safe=False)


# CREATE THE JSON DATA FOR SUB COMPONENT
@csrf_exempt
def json_subcomponent(request):
    data_subcomponent=list(Subcomponent.objects.values('id','subcomponent_name','component__component_name'))
    return JsonResponse(data_subcomponent, safe=False)

# CREATE THE JSON DATA FOR SURVEY PROJECT
@csrf_exempt
def json_surveyproject(request):
    data_surveyproject=list(SurveyProject.objects.values('id','responsible','title_surv',
                                                         'target_population',
                                                         'start_date',
                                                         'end_date',
                                                         'location_survey',
                                                         'date_creation'))
    return JsonResponse(data_surveyproject, safe=False)

# CREATE THE JSON DATA FOR SURVEY DATASET
@csrf_exempt
def json_surveydataset(request):
    data_survey=list(SurveyDataset.objects.values('id','surveyProject__responsible',
                                                  'surveyProject__title_surv',
                                                  'surveyProject__start_date',
                                                  'surveyProject__end_date',
                                                  'surveyProject__location_survey',
                                                  'quest_code','question','response_text','response_num',
                                                  'level_1','level_2'))
    return JsonResponse(data_survey, safe=False)


# CREATE THE JSON DATA FOR USER
@csrf_exempt
def json_user(request):
    data_user=list(User.objects.values('username', 'email'))
    return JsonResponse(data_user, safe=False)


 #DISPLAY PAGE OF ERROR 500
def custom_500(request):
    return render(request, 'mytvddata/cover/500.html', status=500)

 #DISPLAY PAGE OF ERROR 404
def handel404(request, exception):
    return render(request, 'mytvddata/cover/404.html')

# DISPLAY PAGE OF ERROR 403
def handel403(request, exception):
    return render(request, 'mytvddata/cover/403.html')

## FOR OTHER API DATASET TO IMPORT INTO THE DATABASE (NONE GHO)

def select_api_fields(request):
    available_fields = []
    sample_data = []
    details={}  
    context={}  
    payload={}
    db_fields = [f.name for f in ApiExtract._meta.get_fields() if f.concrete and not f.auto_created]

    form = SelectUrlForms(request.POST or None)


    if request.method == "POST" and form.is_valid():
        api_url = form.cleaned_data["api_url"]
        action = request.POST.get("action")

        if action == "import":
            message = import_api_data(api_url)
            messages.success(request, message)

        try:
            response = requests.get(api_url, timeout=5)
            response.raise_for_status()
            payload = response.json()
            details = payload.get("value", [])

            if details:
                # Extraire les champs disponibles
                first_item = details[0]
                available_fields = [
                    {"id": key, "field_name": key, "enabled": True}
                    for key in first_item.keys()
                ]

                # Sauvegarde en base (optionnel)
                for key in first_item.keys():
                    ApiFieldConfig.objects.get_or_create(field_name=key)

                # Prendre les 3 premières lignes
                sample_data = details[:3]

                 # Sauvegarde du mapping choisi
                for field in available_fields:
                    mapping_value = request.POST.get(f"mapping_{field['id']}")
                    if mapping_value:
                        ApiFieldMapping.objects.update_or_create(
                            api_field=field["field_name"],
                            defaults={"db_field": mapping_value}
                        )


                # Import des données si demandé
                if action == "import":
                    message = import_api_data(api_url)
                    messages.success(request, f"{message} Cliquez ci-dessous pour voir les données importées.")

        except Exception as e:
            available_fields = [{"id": "error", "field_name": f"Erreur: {e}", "enabled": False}]
            messages.error(request, f"Erreur lors de la récupération : {e}")

    context = {
        "form": form,
        "available_fields": available_fields,
        "sample_data": sample_data,
        "db_fields": db_fields,
    }

    return render(request, "mytvddata/cover/select_fields.html", context)



def import_api_data(api_url):
    response = requests.get(api_url, timeout=10)
    response.raise_for_status()
    payload = response.json()
    details = payload.get("value", [])

    if not details:
        return "Aucune donnée trouvée."

    # Charger le mapping API → DB
    mappings = {m.api_field: m.db_field for m in ApiFieldMapping.objects.all()}

    # Parcourir les lignes de l’API
    for item in details:
        record_data = {}
        for api_field, db_field in mappings.items():
            if db_field and api_field in item:
                record_data[db_field] = item[api_field]

        # Créer ou mettre à jour l’enregistrement
        ApiExtract.objects.create(**record_data)

    return f"{len(details)} lignes importées avec succès."



def view_imported_data(request):
    records = ApiExtract.objects.all()
    return render(request, "mytvddata/cover/imported_data.html", {
        "records": records
    })


def load_api_other(request):
    api_url = request.session.get("api_url")  # récupère l’URL saisie par l’utilisateur
    if not api_url:
        messages.error(request, "Veuillez d’abord saisir une URL API.")
        return redirect("select_api_fields")

    headers = {"Content-Type": "application/json"}
    allowed_fields = list(ApiFieldConfig.objects.filter(enabled=True).values_list("field_name", flat=True))

    try:
        response = requests.get(api_url, headers=headers, verify=False)
        response.raise_for_status()
        payload = response.json().get("value", [])

        for item in payload:
            filtered_data = {f: item.get(f) for f in allowed_fields}

            ApiExtract.objects.update_or_create(
                api_id=item.get("Id"),
                defaults={
                    "indicator_code": filtered_data.get("IndicatorCode"),
                    "country_code": filtered_data.get("SpatialDim"),
                    "parent_location": filtered_data.get("ParentLocation"),
                    "time_dim": filtered_data.get("TimeDim"),
                    "value": filtered_data.get("Value"),
                    "numeric_value": filtered_data.get("NumericValue"),
                    "date": filtered_data.get("Date"),
                }
            )
        messages.success(request, "API data saved successfully!")

    except Exception as e:
        messages.error(request, f"Error fetching API: {e}")

    data_from_db = ApiExtract.objects.all().order_by("-date")
    return render(request, "mytvddata/cover/api_data_other.html", {"api_data": data_from_db})

## FUnction for StaffMaping

### GOOD -------------------------# Liste principale avec formulaire vide pour le modal "create"

def staff_list(request):
    staff_qs = StaffMember.objects.all().prefetch_related('diseases', 'country')

    # Filtres GET
    country_filter = request.GET.get("country")
    language_filter = request.GET.get("language")
    disease_filter = request.GET.get("disease")

    if country_filter:
        staff_qs = staff_qs.filter(country__name__icontains=country_filter)

    if language_filter:
        staff_qs = staff_qs.filter(language__icontains=language_filter)  # JSONField ou TextField

    if disease_filter:
        staff_qs = staff_qs.filter(diseases__subcomponent_name__icontains=disease_filter).distinct()

    # Pagination
    paginator = Paginator(staff_qs, 10)  # 10 éléments par page
    page_number = request.GET.get("page")
    staff_page = paginator.get_page(page_number)

    # Distincts pour statistiques
    all_countries = [c.name for s in StaffMember.objects.all() for c in s.country.all()]
    distinct_countries = set(all_countries)

    all_languages = [l for s in StaffMember.objects.all() if s.language for l in s.language]
    distinct_languages = set(all_languages)

    all_diseases = [d.subcomponent_name for s in StaffMember.objects.all() for d in s.diseases.all()]
    distinct_diseases = set(all_diseases)

    form = StaffMemberForm()

    return render(
        request,
        "mytvddata/pages/staff/staff_list.html",
        {
            "staff": staff_page,  # ⚠️ on passe l’objet paginé
            "form": form,
            "countries": distinct_countries,
            "languages": distinct_languages,
            "diseases": distinct_diseases,
            "staff_qs":staff_qs
        }
    )




def staff_create(request):
    if request.method == "POST":
        form = StaffMemberForm(request.POST)
        if form.is_valid():
            staff = form.save()
            send_mail("Nouveau Staff ajouté", f"{staff.name} a été ajouté.", "admin@server.com", [staff.email], fail_silently=True)
            messages.success(request, "Staff ajouté avec succès.")
            return HttpResponse("OK")
        return render(request, "mytvddata/pages/staff/partial_staff_form.html", {"form": form})
    else:
        form = StaffMemberForm()
        return render(request, "mytvddata/pages/staff/partial_staff_form.html", {"form": form})

def staff_update(request, pk):
    staff = get_object_or_404(StaffMember, pk=pk)
    if request.method == "POST":
        form = StaffMemberForm(request.POST, instance=staff)
        if form.is_valid():
            staff = form.save()
            send_mail("Mise à jour Staff", f"{staff.name} a été mis à jour.", "admin@server.com", [staff.email], fail_silently=True)
            messages.success(request, "Staff mis à jour avec succès.")
            return HttpResponse("OK")
        return render(request, "mytvddata/pages/staff/partial_staff_form.html", {"form": form, "staff": staff})
    else:
        form = StaffMemberForm(instance=staff)
        return render(request, "mytvddata/pages/staff/partial_staff_form.html", {"form": form, "staff": staff})

def staff_delete(request, pk):
    staff = get_object_or_404(StaffMember, pk=pk)
    if request.method == "POST":
        staff.delete()
        messages.success(request, "Staff supprimé.")
    return redirect("staff_list")




def staff_export_xlsx(request):
    staff = StaffMember.objects.all().prefetch_related('diseases', 'country')

    # ✅ Statistiques globales
    all_countries = [c.name for s in staff for c in s.country.all()]
    distinct_countries = set(all_countries)
    all_diseases = [d.subcomponent_name for s in staff for d in s.diseases.all()]
    distinct_diseases = set(all_diseases)

    # ✅ Création du fichier Excel
    wb = openpyxl.Workbook()

    # --- Feuille 1 : Liste du staff ---
    ws1 = wb.active
    ws1.title = "Staff List"

    header_fill = PatternFill(start_color="0033A0", end_color="0033A0", fill_type="solid")  # Bleu USA
    alt_fill = PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid")    # Gris clair UE
    bold_font = Font(bold=True, size=12, color="FFFFFF")

    headers = [
        "Name", "Email", "Countries", "Position", "Grade", "Telephone",
        "Office Affiliation", "Responsibility", "Languages", "Diseases", "Level Geo"
    ]
    ws1.append(headers)

    # Style header
    for col in range(1, len(headers)+1):
        cell = ws1.cell(row=1, column=col)
        cell.fill = header_fill
        cell.font = bold_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    # Remplir les données avec alternance gris clair
    row_index = 2
    for s in staff:
        countries = ", ".join([c.name for c in s.country.all()])
        diseases = ", ".join([d.subcomponent_name for d in s.diseases.all()])
        languages = ", ".join(s.language) if isinstance(s.language, list) else s.language

        ws1.append([
            s.name,
            getattr(s, "email", ""),
            countries,
            getattr(s, "position", ""),
            getattr(s, "grade", ""),
            getattr(s, "telephone", ""),
            getattr(s, "office_affiliation", ""),
            getattr(s, "responsibility", ""),
            languages,
            diseases,
            getattr(s, "level_geo", "")
        ])

        if row_index % 2 == 0:
            for col in range(1, len(headers)+1):
                ws1.cell(row=row_index, column=col).fill = alt_fill
        row_index += 1

    # --- Feuille 2 : Statistiques ---
    ws2 = wb.create_sheet(title="Statistics")
    ws2["A1"] = "Staff Statistics"
    ws2["A1"].font = Font(bold=True, size=16, color="0033A0")
    ws2.append([])

    stats_data = [
        ("Distinct countries", len(distinct_countries)),
        ("Distinct diseases", len(distinct_diseases)),
        ("Total staff members", staff.count()),
    ]
    for label, value in stats_data:
        ws2.append([label, value])

    # --- Feuille 3 : Graphiques ---
    ws3 = wb.create_sheet(title="Charts")
    ws3["A1"] = "Staff Distribution by Country"
    ws3["A1"].font = Font(bold=True, size=14, color="0033A0")

    # Données pays
    country_counts = {}
    for c in all_countries:
        country_counts[c] = country_counts.get(c, 0) + 1

    ws3.append(["Country", "Count"])
    for country, count in country_counts.items():
        ws3.append([country, count])

    data = Reference(ws3, min_col=2, min_row=2, max_row=ws3.max_row)
    labels = Reference(ws3, min_col=1, min_row=2, max_row=ws3.max_row)

    # ✅ Graphique camembert
    pie = PieChart()
    pie.title = "Staff by Country"
    pie.add_data(data, titles_from_data=False)
    pie.set_categories(labels)
    ws3.add_chart(pie, "E2")

    # ✅ Graphique barres
    bar = BarChart()
    bar.title = "Staff by Country (Bar)"
    bar.add_data(data, titles_from_data=False)
    bar.set_categories(labels)
    bar.y_axis.title = "Staff Count"
    bar.x_axis.title = "Country"
    ws3.add_chart(bar, "E20")

    # ✅ Graphique par maladie
    ws3["A20"] = "Staff Distribution by Disease"
    ws3["A20"].font = Font(bold=True, size=14, color="0033A0")

    ws3.append(["Disease", "Count"])
    disease_counts = {}
    for d in all_diseases:
        disease_counts[d] = disease_counts.get(d, 0) + 1
    for disease, count in disease_counts.items():
        ws3.append([disease, count])

    data_disease = Reference(ws3, min_col=2, min_row=ws3.max_row - len(disease_counts) + 1, max_row=ws3.max_row)
    labels_disease = Reference(ws3, min_col=1, min_row=ws3.max_row - len(disease_counts) + 1, max_row=ws3.max_row)

    pie_disease = PieChart()
    pie_disease.title = "Staff by Disease"
    pie_disease.add_data(data_disease, titles_from_data=False)
    pie_disease.set_categories(labels_disease)
    ws3.add_chart(pie_disease, "E40")

    # ✅ Réponse HTTP avec fichier Excel
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="staff_dashboard.xlsx"'
    wb.save(response)
    return response











def staffs_import_xlsx(request):
    if request.method == "POST" and request.FILES.get("file"):
        excel_file = request.FILES["file"]
        df = pd.read_excel(excel_file)

        for _, row in df.iterrows():
            staff, created = StaffMember.objects.update_or_create(
                email=row["Email"],
                defaults={
                    "name": row["Name"],
                    "email": row.get("Email", ""),
                    "position": row.get("Position", ""),
                    "grade": row.get("Grade", ""),
                    "telephone": row.get("Telephone", ""),
                    "office_affiliation": row.get("Office Affiliation", ""),
                    "responsibility": row.get("Responsibility", ""),
                    "level_geo": row.get("Level Geo", ""),
                }
            )

            # Pays (ManyToMany)
            if "Countries" in row and pd.notna(row["Countries"]):
                countries = [c.strip() for c in row["Countries"].split(",")]
                staff.country.set(Country.objects.filter(name__in=countries))

            # Langues (JSONField)
            if "Languages" in row and pd.notna(row["Languages"]):
                staff.language = [l.strip() for l in row["Languages"].split(",")]

            # Maladies (ManyToMany)
            if "Diseases" in row and pd.notna(row["Diseases"]):
                diseases = [d.strip() for d in row["Diseases"].split(",")]
                staff.diseases.set(Subcomponent.objects.filter(name__in=diseases))

            staff.save()

        messages.success(request, "Importation XLSX réussie.")
    return redirect("staff_list")





@login_required
def staff_import_xlsx(request):
    if request.method == "POST" and request.FILES.get("file"):
        excel_file = request.FILES["file"]
        df = pd.read_excel(excel_file)

        # Step 1: Preview
        headers = list(df.columns)
        preview_data = df.head(20).to_dict(orient="records")

        # Store temporarily in session
        request.session["excel_headers"] = headers
        request.session["excel_data"] = df.to_dict(orient="records")

        return render(request, "mytvddata/pages/staff/staff_import_preview.html", {
            "headers": headers,
            "rows": preview_data,
            "db_fields": [f.name for f in StaffMember._meta.get_fields() if f.concrete],
        })

    return render(request, "mytvddata/pages/staff/staff_import_upload.html")


@login_required
def staff_import_mapping(request):
    if request.method == "POST":
        mapping = request.POST.dict()  # Excel header → DB field mapping
        excel_headers = request.session.get("excel_headers", [])
        excel_data = request.session.get("excel_data", [])

        mapped_records = []
        for row in excel_data:
            record = {}
            for header in excel_headers:
                db_field = mapping.get(header)
                if db_field:
                    record[db_field] = row.get(header)
            mapped_records.append(record)

        # Step 4: Synchronization
        with transaction.atomic():
            for rec in mapped_records:
                staff, created = StaffMember.objects.update_or_create(
                    email=rec.get("email"),
                    defaults=rec
                )

                # Countries (ManyToMany)
                if rec.get("country"):
                    countries = [c.strip() for c in rec["country"].split(",")]
                    staff.country.set(Country.objects.filter(name__in=countries))

                # Languages (JSONField)
                if rec.get("language"):
                    staff.language = [l.strip() for l in rec["language"].split(",")]

                # Diseases (ManyToMany)
                if rec.get("diseases"):
                    diseases = [d.strip() for d in rec["diseases"].split(",")]
                    staff.diseases.set(Subcomponent.objects.filter(name__in=diseases))

                staff.save()

        messages.success(request, "Excel import completed successfully.")
        return redirect("staff_list")

    return render(request, "staff_import_mapping.html", {
        "headers": request.session.get("excel_headers", []),
        "db_fields": [f.name for f in StaffMember._meta.get_fields() if f.concrete],
    })






def staff_template_xlsx(request):
    # Exemple de colonnes attendues
    columns = [
        "Name", "Email", "Position", "Grade",
        "Telephone", "Office Affiliation", "Responsibility",
        "Level Geo", "Countries","Languages", "Diseases"
    ]

    # Exemple de lignes vides ou pré-remplies
    df = pd.DataFrame(columns=columns)
    df.loc[0] = [
         "Jean Dupont", "jean.dupont@example.com", "MOH Staff", "Doctor",
         "+33 123456789", "MOH", "Director",
         "Country", "DRC","fr", "Leprosy" 
    ]
    df.loc[1] = [
         "Maria Sanchez", "maria.sanchez@example.com", "MOH Staff", "Doctor",
         "+34 987654321", "OMS Madrid", "Coordinatrice",
                 "Country", "Burundi","fr", "Yaws"
    ]

    # Réponse HTTP avec fichier Excel
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = 'attachment; filename="staff_template.xlsx"'

    with pd.ExcelWriter(response, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Template")

    return response

## DASHBOARD


def country_list_dashboard(request):
    countries = Country.objects.all().order_by('name')
    return render(request, 'mytvddata/pages/dashboard/country_list_dashboard.html', {'countries': countries})


def country_dashboard(request, pk):
    country = get_object_or_404(Country, pk=pk)
   # indicators = Indicator.objects.filter(country_code=country.cca3).select_related('subcomponent')

     # Résoudre le lien court goo.gl/maps
    resolved_url = country.maps
    if resolved_url and "goo.gl/maps" in resolved_url:
        try:
            response = requests.head(resolved_url, allow_redirects=True)
            resolved_url = response.url   # URL finale Google Maps
        except Exception:
            resolved_url = country.maps   # fallback

    indicators = StoreAPI.objects.filter(country_code=country.cca3).values_list('indicator_code', flat=True)
    api_data = StoreAPI.objects.filter(country_code=country.cca3)
   # api_data = ApiExtract.objects.filter(country_code=country.flag)

       # ✅ Récupération des données (optimisée)
 
    locations = LocationCountry.objects.filter(iso3=country.cca3).only('iso3','name','latitude','longitude')

    # Coordonnées par défaut
    center_lat, center_lon = 0.0, 20.0

    if locations.exists():
        latitudes = [loc.latitude for loc in locations if loc.latitude]
        longitudes = [loc.longitude for loc in locations if loc.longitude]
        if latitudes and longitudes:
            center_lat = sum(latitudes) / len(latitudes)
            center_lon = sum(longitudes) / len(longitudes)

    m = folium.Map(
        location=[center_lat, center_lon],
        zoom_start=4,
        tiles="cartodb positron",
        scrollWheelZoom=True
    )


    folium.LayerControl().add_to(m)
  
    # ✅ Ajout des markers
    for location in locations:
        if location.latitude and location.longitude:
            folium.Marker(
                location=(location.latitude, location.longitude),
                popup=f"<b>{location.name}</b><br>Lat: {location.latitude}, Lon: {location.longitude}",
                tooltip=location.name,
                icon=folium.Icon(color="blue", icon="globe", prefix="fa")  # icône globe
            ).add_to(m)
 
     # ✅ Conversion en HTML (IMPORTANT : en dehors de la boucle)
    map_html = m._repr_html_()


    ## Récupérer toutes les maladies (subcomponents) liées aux indicateurs du pays
    # ... ton code pour la map ...

    subcomponents = Subcomponent.objects.filter(
        subcomponent_indicator__indicator_code__in=StoreAPI.objects.filter(
            country_code=country.cca3
        ).values_list("indicator_code", flat=True)
    ).distinct()

    dashboard_data = []
    all_years = set()

    for sub in subcomponents:
        indicators = Indicator.objects.filter(subcomponent=sub)

        indicator_results = []
        for ind in indicators:
            results = StoreAPI.objects.filter(
                country_code=country.cca3,
                indicator_code=ind.indicator_code
            ).order_by("time_dim")

            # ⚠️ Ne garder que les indicateurs avec au moins une valeur
            if results.exists():
                # Récupérer min et max year
                min_year = results.first().time_dim
                max_year = results.last().time_dim

                variation = None
                # ✅ Cas quantitatif
                if ind.type_indicator == "Quanti_ind":
                    min_val = results.first().numeric_value
                    max_val = results.last().numeric_value
                    if min_val is not None and max_val is not None and min_val != 0:
                        try:
                            variation = ((max_val - min_val) / min_val) * 100
                        except (ValueError, TypeError, ZeroDivisionError):
                            variation = None

                # ✅ Cas qualitatif : pas de variation
                elif ind.type_indicator == "Quali_ind":
                    min_val = results.first().alpha_value
                    max_val = results.last().alpha_value
                    variation = None

                indicator_results.append({
                    "indicator": ind,
                    "results": results,
                    "min_year": min_year,
                    "max_year": max_year,
                    "variation": variation,
                    "target": ind.indicator_target,             # champ Target
                    "performance": ind.performance_indicator,   # champ Performance
                    "category": ind.category_indicator,         # champ Category
                    "data_source": ind.indicator_source,         # champ Data Source
                    "type_indicator": ind.type_indicator
                })


            # Collecter les années pour all_years
            for res in results:
                if res.time_dim is not None:
                    try:
                        all_years.add(int(res.time_dim))
                    except (ValueError, TypeError):
                        pass

        dashboard_data.append({
            "subcomponent": sub,
            "indicators": indicator_results,
        })

    # Transformer en liste triée
    all_years = sorted(all_years)

    context = {
        "country": country,
        "api_data": StoreAPI.objects.filter(country_code=country.cca3),
        "resolved_url": resolved_url,
        "map_html": map_html,
        "dashboard_data": dashboard_data,
        "all_years": all_years,
        "api_data": api_data,
    }
   
    return render(request, 'mytvddata/pages/dashboard/country_profile.html', context)






def export_country_word(request, pk):
    country = get_object_or_404(Country, pk=pk)

    api_data = StoreAPI.objects.filter(country_code=country.cca3)

    indicators = Indicator.objects.filter(indicator_code__in=api_data.values_list("indicator_code", flat=True))
    indicator_map = {ind.indicator_code: ind for ind in indicators}

    doc = Document()
    doc.add_heading(f"Profil du pays : {country.name}", level=1)
    doc.add_paragraph(f"Capitale : {country.capital}")
    doc.add_paragraph(f"Subregion : {country.subregion}")
    doc.add_paragraph(f"Area : {country.area}")
    doc.add_paragraph(f"Population : {country.population}")
    doc.add_paragraph(f"Langues : {country.languages}")
    doc.add_paragraph(f"Class : {country.country_class}")

    # --- Indicateurs quantitatifs ---
    quanti_codes = [code for code, ind in indicator_map.items() if ind.type_indicator == "Quanti_ind"]
    if quanti_codes:
        para = doc.add_paragraph()
        run = para.add_run("Indicateurs quantitatifs (Quanti_ind)")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 128, 0)  # Vert

        table = doc.add_table(rows=1, cols=7)
        table.style = "Light List Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Maladie"
        hdr_cells[1].text = "Indicateur"
        hdr_cells[2].text = "Année min"
        hdr_cells[3].text = "Année max"
        hdr_cells[4].text = "Valeur min"
        hdr_cells[5].text = "Valeur max"
        hdr_cells[6].text = "Variation (%)"

        for code in quanti_codes:
            results = api_data.filter(indicator_code=code).order_by("time_dim")
            if results.exists():
                ind = indicator_map[code]
                min_val = results.first().numeric_value
                max_val = results.last().numeric_value
                min_year = results.first().time_dim
                max_year = results.last().time_dim

                variation = None
                if min_val and max_val and min_val != 0:
                    try:
                        variation = ((max_val - min_val) / min_val) * 100
                    except Exception:
                        variation = None

                row_cells = table.add_row().cells
                # Maladie en bleu
                run = row_cells[0].paragraphs[0].add_run(ind.subcomponent.subcomponent_name if ind.subcomponent else "-")
                run.font.color.rgb = RGBColor(0, 0, 255)
                # Indicateur en gras
                run = row_cells[1].paragraphs[0].add_run(ind.indicator_name)
                run.bold = True
                row_cells[2].text = str(min_year)
                row_cells[3].text = str(max_year)
                row_cells[4].text = str(min_val) if min_val is not None else "-"
                row_cells[5].text = str(max_val) if max_val is not None else "-"

                if variation is not None:
                    para = row_cells[6].paragraphs[0]
                    if ind.performance_indicator == "Low" and ind.category_indicator in ["prevalence","Death","incidence"] and variation < 0:
                        run = para.add_run(f"↓ {variation:.2f} %")
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif ind.performance_indicator == "Hight" and ind.category_indicator == "strategy" and variation > 0:
                        run = para.add_run(f"↑ {variation:.2f} %")
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    else:
                        run = para.add_run(f"↓ {variation:.2f} %")
                        run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    row_cells[6].text = "–"

    # --- Indicateurs qualitatifs ---
    quali_codes = [code for code, ind in indicator_map.items() if ind.type_indicator == "Quali_ind"]
    if quali_codes:
        para = doc.add_paragraph()
        run = para.add_run("Indicateurs qualitatifs (Quali_ind)")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 140, 0)  # Orange

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light List Accent 2"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Maladie"
        hdr_cells[1].text = "Indicateur"
        hdr_cells[2].text = "Année"
        hdr_cells[3].text = "Valeur"
        hdr_cells[4].text = "Source"

        for code in quali_codes:
            results = api_data.filter(indicator_code=code).order_by("time_dim")
            ind = indicator_map[code]
            for data in results:
                row_cells = table.add_row().cells
                run = row_cells[0].paragraphs[0].add_run(ind.subcomponent.subcomponent_name if ind.subcomponent else "-")
                run.font.color.rgb = RGBColor(0, 0, 255)  # Bleu
                run = row_cells[1].paragraphs[0].add_run(ind.indicator_name)
                run.bold = True
                row_cells[2].text = str(data.time_dim)
                row_cells[3].text = str(data.alpha_value) if data.alpha_value else "-"
                row_cells[4].text = ind.indicator_source or "-"

    # --- Légende ---
    doc.add_page_break()
    doc.add_heading("Légende", level=2)
    legend_table = doc.add_table(rows=1, cols=2)
    legend_table.style = "Light List Accent 3"
    hdr = legend_table.rows[0].cells
    hdr[0].text = "Symbole / Couleur"
    hdr[1].text = "Signification"

    rows = [
        ("↑ Vert", "Amélioration souhaitée"),
        ("↓ Rouge", "Détérioration"),
        ("– Gris", "Variation nulle"),
        ("–", "Données non disponibles"),
        ("Quanti_ind", "Valeurs numériques par année, variation calculée"),
        ("Quali_ind", "Valeurs alphabétiques (alpha_value), pas de variation calculée"),
    ]

    for symbol, meaning in rows:
        row = legend_table.add_row().cells
        run = row[0].paragraphs[0].add_run(symbol)
        if "Vert" in symbol:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif "Rouge" in symbol:
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif "Gris" in symbol:
            run.font.color.rgb = RGBColor(128, 128, 128)
        row[1].text = meaning

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="{country.name}_profile.docx"'
    doc.save(response)
    return response





def export_country_dashboard_excel(request, pk):
    country = get_object_or_404(Country, pk=pk)

    # Récupération des subcomponents et indicateurs comme dans country_dashboard
    subcomponents = Subcomponent.objects.filter(
        subcomponent_indicator__indicator_code__in=StoreAPI.objects.filter(
            country_code=country.cca3
        ).values_list("indicator_code", flat=True)
    ).distinct()

    dashboard_data = []
    all_years = set()

    for sub in subcomponents:
        indicators = Indicator.objects.filter(subcomponent=sub)
        indicator_results = []
        for ind in indicators:
            results = StoreAPI.objects.filter(
                country_code=country.cca3,
                indicator_code=ind.indicator_code
            ).order_by("time_dim")

            if results.exists():
                min_year = results.first().time_dim
                max_year = results.last().time_dim
                variation = None
                if ind.type_indicator == "Quanti_ind":
                    min_val = results.first().numeric_value
                    max_val = results.last().numeric_value
                    if min_val and max_val and min_val != 0:
                        variation = ((max_val - min_val) / min_val) * 100
                indicator_results.append({
                    "indicator": ind,
                    "results": results,
                    "variation": variation,
                    "target": ind.indicator_target,
                    "performance": ind.performance_indicator,
                    "category": ind.category_indicator,
                    "data_source": ind.indicator_source,
                    "type_indicator": ind.type_indicator
                })
                for res in results:
                    if res.time_dim:
                        try:
                            all_years.add(int(res.time_dim))
                        except:
                            pass
        dashboard_data.append({"subcomponent": sub, "indicators": indicator_results})

    all_years = sorted(all_years)

    # ✅ Création du fichier Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"{country.name} Dashboard"

    # Styles USA + UE
    header_fill = PatternFill(start_color="0033A0", end_color="0033A0", fill_type="solid")  # Bleu USA
    eu_fill = PatternFill(start_color="FFD700", end_color="FFD700", fill_type="solid")      # Jaune UE
    bold_font = Font(bold=True, size=12, color="FFFFFF")
    normal_font = Font(size=11)

    # ✅ Profil du pays
    ws["A1"] = f"Country Dashboard: {country.name} ({country.cca3})"
    ws["A1"].font = Font(bold=True, size=16, color="0033A0")
    ws.append([])
    profile_data = [
        ("Official name", country.official),
        ("Capital", country.capital),
        ("Subregion", country.subregion),
        ("Area (km²)", country.area),
        ("Population", country.population),
        ("Languages", country.languages),
        ("Class", country.country_class),
    ]
    for label, value in profile_data:
        ws.append([label, value])
    ws.append([])

    # ✅ Indicateurs par sous-composant
    for section in dashboard_data:
        ws.append([section["subcomponent"].subcomponent_name])
        header = ["Indicator", "Target"] + [str(y) for y in all_years] + ["Variation (%)", "Source"]
        ws.append(header)
        # Style header
        for col in range(1, len(header)+1):
            cell = ws.cell(row=ws.max_row, column=col)
            cell.fill = header_fill
            cell.font = bold_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

        for ind in section["indicators"]:
            row = [ind["indicator"].indicator_name, ind["target"]]
            for year in all_years:
                res = ind["results"].filter(time_dim=year).first()
                if res:
                    if ind["type_indicator"] == "Quanti_ind":
                        row.append(res.numeric_value if res.numeric_value is not None else "-")
                    else:
                        row.append(res.alpha_value if res.alpha_value else "-")
                else:
                    row.append("-")
            if ind["type_indicator"] == "Quanti_ind":
                row.append(f"{ind['variation']:.2f}%" if ind["variation"] is not None else "–")
            else:
                row.append("–")
            row.append(ind["data_source"])
            ws.append(row)
        ws.append([])

    # ✅ Réponse HTTP avec fichier Excel
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="country_dashboard_{country.cca3}.xlsx"'
    wb.save(response)
    return response




def preview_json(request):
    """
    Vue pour prévisualiser le fichier JSON uploadé et proposer l'import en base.
    """
    preview_data = None

    if request.method == "POST" and request.FILES.get("json_file"):
        json_file = request.FILES["json_file"]
        try:
            preview_data = json.load(json_file)  # Charger le contenu JSON
            request.session["preview_data"] = preview_data  # stocker en session
            messages.success(request, "Fichier JSON chargé avec succès. Voici la prévisualisation.")
        except Exception as e:
            messages.error(request, f"Erreur lors du chargement du fichier JSON : {e}")

    return render(request, "mytvddata/pages/country/preview_json.html", {"preview_data": preview_data})


def import_json(request):
    """
    Vue pour importer en base le JSON prévisualisé.
    """
    preview_data = request.session.get("preview_data")

    if preview_data:
        try:
            for c in preview_data:
                Country.objects.update_or_create(
                    cca3=c.get("cca3"),
                    defaults={
                        "flag": c.get("flag"),
                        "name": c.get("name"),
                        "official": c.get("official"),
                        "capital": ", ".join(eval(c.get("capital"))) if c.get("capital") else "",
                        "subregion": c.get("subregion"),
                        "area": c.get("area"),
                        "population": c.get("population"),
                        "languages": c.get("languages"),
                        "maps": c.get("data_source"),  # provisoire
                        "ref_data": c.get("ref_data"),
                        "data_source": c.get("data_source"),
                        "country_class": c.get("country_class"),
                    }
                )
            messages.success(request, "Importation réussie ! Les pays ont été ajoutés à la base.")
        except Exception as e:
            messages.error(request, f"Erreur lors de l’importation : {e}")
    else:
        messages.error(request, "Aucune donnée JSON à importer.")

    return redirect("preview_json")




######NEW SURVEY PROJECT


def upload_datasurvey(request):
    if request.method == "POST":
        form = SurveyUploadForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES["file"]

            # Création du projet Survey
            project = SurveyProject.objects.create(
                responsible=form.cleaned_data.get("responsible"),
                title_surv=form.cleaned_data.get("title_surv"),
                target_population=form.cleaned_data.get("target_population"),
                start_date=form.cleaned_data.get("start_date"),
                end_date=form.cleaned_data.get("end_date"),
                location_survey=form.cleaned_data.get("location_survey"),
            )

            # Lecture du fichier Excel avec pandas
            df = pd.read_excel(excel_file)

            # Boucle d’insertion des données
            for _, row in df.iterrows():
                SurveyDataset.objects.create(
                    surveyProject=project,
                    quest_code=row.get("quest_code"),
                    question=row.get("question"),
                    response_text=row.get("response_text"),
                    response_num=row.get("response_num"),
                    level_1=row.get("level_1"),
                    level_2=row.get("level_2"),
                )

            messages.success(request, "Importation réussie !")
            return redirect("survey_datalist")  # 🔥 corriger ici
    else:
        form = SurveyUploadForm()

    return render(request, "mytvddata/pages/survey/upload_datasurvey.html", {"form": form})


def survey_datalist(request):
    projects = SurveyProject.objects.prefetch_related("project_surveyData").all()
    return render(request, "mytvddata/pages/survey/list_survey.html", {"projects": projects})


# COUNTRY LOACTION: LATITUDE AND LONGITUDE






# ✅ Liste des pays (CRUD - Read)
def geocountry_list(request):
    query = request.GET.get("q", "")
    countries = LocationCountry.objects.all()
    if query:
        countries = countries.filter(name__icontains=query)

    paginator = Paginator(countries, 10)
    page_number = request.GET.get("page")
    page_obj = paginator.get_page(page_number)

    # Si Ajax → renvoyer juste le fragment
    if request.headers.get("x-requested-with") == "XMLHttpRequest":
        return render(
            request,
            "mytvddata/pages/locationCountry/_country_table.html",
            {"countries": page_obj, "query": query}
        )

    return render(
        request,
        "mytvddata/pages/locationCountry/locationCountry_list.html",
        {"countries": page_obj, "query": query}
    )


# ✅ Importer les données JSON
def geocountry_import_json(request):
    if request.method == "POST" and request.FILES.get("json_file"):
        try:
            json_file = request.FILES["json_file"]
            data = json.load(json_file)
            request.session["json_preview"] = data
            messages.success(request, "Prévisualisation des données chargée.")
            return redirect("geocountry_preview")
        except json.JSONDecodeError as e:
            messages.error(request, f"Erreur JSON : {e}")
            return redirect("geocountry_import_json")

    return render(request, "mytvddata/pages/locationCountry/geoimport_json.html")

# ✅ Prévisualiser les données importées
def geocountry_preview(request):
    data = request.session.get("json_preview", [])
    return render(request, "mytvddata/pages/locationCountry/geocountry_preview.html", {"data": data})

# ✅ Synchroniser avec la base SQL Server
def geocountry_sync_to_db(request):
    data = request.session.get("json_preview", [])
    if not data:
        messages.error(request, "Aucune donnée à synchroniser.")
        return redirect("geocountry_import_json")

    for item in data:
        LocationCountry.objects.update_or_create(
            iso3=item.get("CountryName")[:3].upper(),
            defaults={
                "name": item.get("CountryName"),
                "latitude": item.get("CountryLat"),
                "longitude": item.get("CountryLng"),
            },
        )
    messages.success(request, "Les données ont été synchronisées avec la base SQL Server.")
    return redirect("geocountry_list")

# ✅ CRUD Ajax
def geocountry_create(request):
    if request.method == "POST":
        form = LocationCountryForm(request.POST)
        if form.is_valid():
            form.save()
            return JsonResponse({"success": True})
        return JsonResponse({"success": False, "errors": form.errors})

def geocountry_update(request, pk):
    country = get_object_or_404(LocationCountry, pk=pk)
    if request.method == "POST":
        form = LocationCountryForm(request.POST, instance=country)
        if form.is_valid():
            form.save()
            return JsonResponse({"success": True})
        return JsonResponse({"success": False, "errors": form.errors})

def geocountry_delete(request, pk):
    country = get_object_or_404(LocationCountry, pk=pk)
    if request.method == "POST":
        country.delete()
        return JsonResponse({"success": True})
    return JsonResponse({"success": False})

# ✅ Export JSON
def export_json_geocountry(request):
    countries = LocationCountry.objects.all().values("iso3", "name", "latitude", "longitude")
    data = list(countries)
    response = HttpResponse(json.dumps(data, indent=4), content_type="application/json")
    response["Content-Disposition"] = 'attachment; filename="countries.json"'
    return response

# ✅ Export CSV
def export_csv_geocountry(request):
    countries = LocationCountry.objects.all().values("iso3", "name", "latitude", "longitude")
    response = HttpResponse(content_type="text/csv")
    response["Content-Disposition"] = 'attachment; filename="countries.csv"'
    writer = csv.writer(response)
    writer.writerow(["ISO3", "Name", "Latitude", "Longitude"])
    for c in countries:
        writer.writerow([c["iso3"], c["name"], c["latitude"], c["longitude"]])
    return response


def geocountry_data(request):
    countries = LocationCountry.objects.all().values("id", "iso3", "name", "latitude", "longitude")
    return JsonResponse({"data": list(countries)})




def regional_factsheet_select(request):

    """Page de sélection des maladies (subcomponents)."""

    subcomponents = Subcomponent.objects.all().order_by("subcomponent_name")
    return render(request, "mytvddata/pages/dashboard/factsheet_select.html", {"subcomponents": subcomponents})


def regional_factsheet_results(request, subcomponent_id):
    """Page d'analyse régionale regroupée par type d'indicateur."""
    subcomponent = get_object_or_404(Subcomponent, pk=subcomponent_id)
    indicators = Indicator.objects.filter(subcomponent=subcomponent)

    all_years = sorted(StoreAPI.objects.values_list("time_dim", flat=True).distinct())

    grouped_indicators = {}
    for ind in indicators:
        cat = ind.category_indicator or "Autres"
        grouped_indicators.setdefault(cat, []).append(ind)

    country_summary = []
    for country in Country.objects.all().order_by("name"):
        country_data = StoreAPI.objects.filter(
            country_code=country.cca3,
            indicator_code__in=indicators.values_list("indicator_code", flat=True)
        )
        progress = []
        for ind in indicators:
            values = {v.time_dim: v for v in country_data.filter(indicator_code=ind.indicator_code)}
            if not values:
                continue

            if ind.type_indicator == "Quanti_ind":
                vals = [v.numeric_value for v in values.values() if v.numeric_value is not None]
                if not vals:
                    continue
                variation = None
                if len(vals) >= 2 and vals[0] != 0:
                    variation = ((vals[-1] - vals[0]) / vals[0]) * 100
                progress.append({
                    "indicator": ind,
                    "values": {year: (v.numeric_value if v else None) for year, v in values.items()},
                    "variation": variation
                })

            elif ind.type_indicator == "Quali_ind":
                vals = [v.alpha_value for v in values.values() if v.alpha_value]
                if not vals:
                    continue
                progress.append({
                    "indicator": ind,
                    "values": {year: (v.alpha_value if v else None) for year, v in values.items()},
                    "variation": None
                })

        if progress:
            country_summary.append({
                "country": country,
                "progress": progress,
                "score": sum([p["variation"] or 0 for p in progress if p["variation"] is not None])
            })

    ranked_countries = sorted(country_summary, key=lambda c: c["score"], reverse=False)

    # ✅ Calcul des statistiques globales
    total_countries = len(country_summary)

    # Quali_ind : nombre de pays avec au moins une valeur
    quali_count = sum(
        1 for c in country_summary for p in c["progress"]
        if p["indicator"].type_indicator == "Quali_ind" and any(p["values"].values())
    )

    # Quanti_ind : statut Progrès / Recul / Stable
    progress_count = sum(1 for c in country_summary if c["score"] < 0)
    recul_count = sum(1 for c in country_summary if c["score"] > 0)
    stable_count = sum(1 for c in country_summary if c["score"] == 0)

    stats = {
        "total_countries": total_countries,
        "quali_count": quali_count,
        "quali_percent": (quali_count / total_countries * 100) if total_countries else 0,
        "progress_count": progress_count,
        "progress_percent": (progress_count / total_countries * 100) if total_countries else 0,
        "recul_count": recul_count,
        "recul_percent": (recul_count / total_countries * 100) if total_countries else 0,
        "stable_count": stable_count,
        "stable_percent": (stable_count / total_countries * 100) if total_countries else 0,
    }

    return render(request, "mytvddata/pages/dashboard/factsheet_results.html", {
        "subcomponent": subcomponent,
        "grouped_indicators": grouped_indicators,
        "all_years": all_years,
        "country_summary": country_summary,
        "ranked_countries": ranked_countries,
        "stats": stats
    })

def export_regional_factsheet_word(request, subcomponent_id):
    """Export Word du factsheet régional avec tableau pivoté par années et statistiques globales."""
    subcomponent = get_object_or_404(Subcomponent, pk=subcomponent_id)
    indicators = Indicator.objects.filter(subcomponent=subcomponent)
    api_data = StoreAPI.objects.filter(indicator_code__in=indicators.values_list("indicator_code", flat=True))

    all_years = sorted(StoreAPI.objects.values_list("time_dim", flat=True).distinct())

    grouped_indicators = {}
    for ind in indicators:
        cat = ind.category_indicator or "Autres"
        grouped_indicators.setdefault(cat, []).append(ind)

    # Construire les données par pays
    country_summary = []
    for country in Country.objects.all().order_by("name"):
        country_data = api_data.filter(country_code=country.cca3)
        progress = []
        for ind in indicators:
            values = {v.time_dim: v for v in country_data.filter(indicator_code=ind.indicator_code)}
            if not values:
                continue

            if ind.type_indicator == "Quanti_ind":
                vals = [v.numeric_value for v in values.values() if v.numeric_value is not None]
                if not vals:
                    continue
                variation = None
                if len(vals) >= 2 and vals[0] != 0:
                    variation = ((vals[-1] - vals[0]) / vals[0]) * 100
                progress.append({
                    "indicator": ind,
                    "values": {year: (v.numeric_value if v else None) for year, v in values.items()},
                    "variation": variation
                })
            elif ind.type_indicator == "Quali_ind":
                vals = [v.alpha_value for v in values.values() if v.alpha_value]
                if not vals:
                    continue
                progress.append({
                    "indicator": ind,
                    "values": {year: (v.alpha_value if v else None) for year, v in values.items()},
                    "variation": None
                })

        if progress:
            country_summary.append({
                "country": country,
                "progress": progress,
                "score": sum([p["variation"] or 0 for p in progress if p["variation"] is not None])
            })

    ranked_countries = sorted(country_summary, key=lambda c: c["score"], reverse=False)

    # ✅ Calcul des statistiques globales
    total_countries = len(country_summary)
    quali_count = sum(
        1 for c in country_summary for p in c["progress"]
        if p["indicator"].type_indicator == "Quali_ind" and any(p["values"].values())
    )
    progress_count = sum(1 for c in country_summary if c["score"] < 0)
    recul_count = sum(1 for c in country_summary if c["score"] > 0)
    stable_count = sum(1 for c in country_summary if c["score"] == 0)

    doc = Document()
    doc.add_heading(f"Factsheet régional : {subcomponent.subcomponent_name}", level=1)

    # ✅ Ajouter les statistiques globales en introduction
    doc.add_heading("Statistiques globales", level=2)
    doc.add_paragraph(f"Total des pays analysés : {total_countries}")
    doc.add_paragraph(f"Pays avec indicateurs qualitatifs : {quali_count} ({(quali_count/total_countries*100 if total_countries else 0):.2f}%)")
    doc.add_paragraph(f"Pays en progrès : {progress_count} ({(progress_count/total_countries*100 if total_countries else 0):.2f}%)")
    doc.add_paragraph(f"Pays en recul : {recul_count} ({(recul_count/total_countries*100 if total_countries else 0):.2f}%)")
    doc.add_paragraph(f"Pays stables : {stable_count} ({(stable_count/total_countries*100 if total_countries else 0):.2f}%)")

    # ✅ Boucle par catégorie avec tableau pivoté
    for category, inds in grouped_indicators.items():
        doc.add_heading(category.title(), level=2)
        table = doc.add_table(rows=1, cols=2 + len(all_years) + 1)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Pays"
        hdr[1].text = "Indicateur"
        for i, year in enumerate(all_years):
            hdr[2+i].text = str(year)
        hdr[-1].text = "Variation (%)"

        for country in Country.objects.all().order_by("name"):
            country_data = api_data.filter(country_code=country.cca3)
            for ind in inds:
                values = {v.time_dim: v for v in country_data.filter(indicator_code=ind.indicator_code)}
                if not values:
                    continue

                row = table.add_row().cells
                row[0].text = country.name
                row[1].text = ind.indicator_name

                for i, year in enumerate(all_years):
                    val = values.get(year)
                    if val:
                        if ind.type_indicator == "Quanti_ind":
                            row[2+i].text = str(val.numeric_value) if val.numeric_value is not None else "-"
                        else:
                            row[2+i].text = str(val.alpha_value) if val.alpha_value else "-"
                    else:
                        row[2+i].text = "-"

                variation = None
                if ind.type_indicator == "Quanti_ind":
                    vals = [v.numeric_value for v in values.values() if v.numeric_value is not None]
                    if len(vals) >= 2 and vals[0] != 0:
                        variation = ((vals[-1] - vals[0]) / vals[0]) * 100

                if variation is not None:
                    run = row[-1].paragraphs[0].add_run(f"{variation:.2f} %")
                    if variation < 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    else:
                        run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    row[-1].text = "–"

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="factsheet_regional_{subcomponent.subcomponent_name}.docx"'
    doc.save(response)
    return response



# TO EXPORT COUNTRY PROFILE IN EXCEL FORMAT






def export_regional_factsheet_excel(request, subcomponent_id):
    subcomponent = get_object_or_404(Subcomponent, pk=subcomponent_id)
    indicators = Indicator.objects.filter(subcomponent=subcomponent)
    api_data = StoreAPI.objects.filter(indicator_code__in=indicators.values_list("indicator_code", flat=True))
    all_years = sorted(StoreAPI.objects.values_list("time_dim", flat=True).distinct())

    # Regrouper les indicateurs par catégorie
    grouped_indicators = {}
    for ind in indicators:
        cat = ind.category_indicator or "Autres"
        grouped_indicators.setdefault(cat, []).append(ind)

    # Construire les données par pays
    country_summary = []
    for country in Country.objects.all().order_by("name"):
        country_data = api_data.filter(country_code=country.cca3)
        progress = []
        for ind in indicators:
            values = {v.time_dim: v for v in country_data.filter(indicator_code=ind.indicator_code)}
            if not values:
                continue
            if ind.type_indicator == "Quanti_ind":
                vals = [v.numeric_value for v in values.values() if v.numeric_value is not None]
                if not vals:
                    continue
                variation = None
                if len(vals) >= 2 and vals[0] != 0:
                    variation = ((vals[-1] - vals[0]) / vals[0]) * 100
                progress.append({
                    "indicator": ind,
                    "values": {year: (v.numeric_value if v else None) for year, v in values.items()},
                    "variation": variation
                })
            elif ind.type_indicator == "Quali_ind":
                vals = [v.alpha_value for v in values.values() if v.alpha_value]
                if not vals:
                    continue
                progress.append({
                    "indicator": ind,
                    "values": {year: (v.alpha_value if v else None) for year, v in values.items()},
                    "variation": None
                })
        if progress:
            country_summary.append({
                "country": country,
                "progress": progress,
                "score": sum([p["variation"] or 0 for p in progress if p["variation"] is not None])
            })

    # ✅ Statistiques globales
    total_countries = len(country_summary)
    quali_count = sum(
        1 for c in country_summary for p in c["progress"]
        if p["indicator"].type_indicator == "Quali_ind" and any(p["values"].values())
    )
    progress_count = sum(1 for c in country_summary if c["score"] < 0)
    recul_count = sum(1 for c in country_summary if c["score"] > 0)
    stable_count = sum(1 for c in country_summary if c["score"] == 0)

    # ✅ Création du fichier Excel
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Factsheet {subcomponent.subcomponent_name}"

    # Style USA + UE
    header_fill = PatternFill(start_color="0033A0", end_color="0033A0", fill_type="solid")  # Bleu USA
    eu_fill = PatternFill(start_color="FFD700", end_color="FFD700", fill_type="solid")      # Jaune UE
    bold_font = Font(bold=True, size=12, color="FFFFFF")
    normal_font = Font(size=11)

    # ✅ Titre
    ws["A1"] = f"Regional Factsheet: {subcomponent.subcomponent_name}"
    ws["A1"].font = Font(bold=True, size=16, color="0033A0")
    ws.append([])

    # ✅ Statistiques globales
    stats_data = [
        ("Total countries", total_countries),
        ("Countries with qualitative indicators", f"{quali_count} ({(quali_count/total_countries*100 if total_countries else 0):.2f}%)"),
        ("Countries in progress", f"{progress_count} ({(progress_count/total_countries*100 if total_countries else 0):.2f}%)"),
        ("Countries in decline", f"{recul_count} ({(recul_count/total_countries*100 if total_countries else 0):.2f}%)"),
        ("Stable countries", f"{stable_count} ({(stable_count/total_countries*100 if total_countries else 0):.2f}%)"),
    ]
    for label, value in stats_data:
        ws.append([label, value])

    ws.append([])

    # ✅ Tableaux par catégorie
    for category, inds in grouped_indicators.items():
        ws.append([category])
        header = ["Country", "Indicator"] + [str(y) for y in all_years] + ["Variation (%)"]
        ws.append(header)
        # Style header
        for col in range(1, len(header)+1):
            cell = ws.cell(row=ws.max_row, column=col)
            cell.fill = header_fill
            cell.font = bold_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

        for country in Country.objects.all().order_by("name"):
            country_data = api_data.filter(country_code=country.cca3)
            for ind in inds:
                values = {v.time_dim: v for v in country_data.filter(indicator_code=ind.indicator_code)}
                if not values:
                    continue
                row = [country.name, ind.indicator_name]
                for year in all_years:
                    val = values.get(year)
                    if val:
                        if ind.type_indicator == "Quanti_ind":
                            row.append(val.numeric_value if val.numeric_value is not None else "-")
                        else:
                            row.append(val.alpha_value if val.alpha_value else "-")
                    else:
                        row.append("-")
                if ind.type_indicator == "Quanti_ind":
                    vals = [v.numeric_value for v in values.values() if v.numeric_value is not None]
                    if len(vals) >= 2 and vals[0] != 0:
                        variation = ((vals[-1] - vals[0]) / vals[0]) * 100
                        row.append(f"{variation:.2f}%")
                    else:
                        row.append("–")
                else:
                    row.append("–")
                ws.append(row)
        ws.append([])

    # ✅ Réponse HTTP avec fichier Excel
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="factsheet_regional_{subcomponent.subcomponent_name}.xlsx"'
    wb.save(response)
    return response
